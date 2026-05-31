"""
Academic Research Writeup Agent — 5-Chapter Format
Generates a complete 5-chapter academic research project in Microsoft Word (.docx).
Supports undergraduate and postgraduate research levels.
"""

import os
import sys
import re
import anthropic
import tempfile

# Visualization imports (optional — graceful fallback if not installed)
try:
    import matplotlib
    matplotlib.use('Agg')  # Non-interactive backend
    import matplotlib.pyplot as plt
    import numpy as np
    VISUALIZATION_AVAILABLE = True
except ImportError:
    VISUALIZATION_AVAILABLE = False

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

import config

# ─────────────────────────────────────────────────────────
#  CHAPTER SELECTION HELPERS
# ─────────────────────────────────────────────────────────

def parse_chapters(chapters_input) -> list:
    """
    Convert any chapters input into a sorted list of chapter numbers (1–5).

    Accepted formats:
      int  5        → [1,2,3,4,5]   integer = "up to N" (backward-compat)
      "all"         → [1,2,3,4,5]
      "3"           → [3]            single chapter
      "3-5"         → [3,4,5]        range
      "1,3,5"       → [1,3,5]        comma list
      "1,3-5"       → [1,3,4,5]      mixed
      [3,4,5]       → [3,4,5]        already a list
    """
    if chapters_input is None:
        return list(range(1, 6))

    if isinstance(chapters_input, list):
        nums = sorted({int(x) for x in chapters_input if 1 <= int(x) <= 5})
        return nums or list(range(1, 6))

    # Plain integer → backward-compat "up to N"
    if isinstance(chapters_input, (int, float)):
        n = max(1, min(5, int(chapters_input)))
        return list(range(1, n + 1))

    s = str(chapters_input).strip().lower()
    if s in ("all", "1-5"):
        return list(range(1, 6))

    result = set()
    for part in s.split(","):
        part = part.strip()
        if not part:
            continue
        if "-" in part:
            sides = part.split("-", 1)
            try:
                a, b = int(sides[0].strip()), int(sides[1].strip())
                for n in range(min(a, b), max(a, b) + 1):
                    if 1 <= n <= 5:
                        result.add(n)
            except ValueError:
                pass
        else:
            try:
                n = int(part)
                if 1 <= n <= 5:
                    result.add(n)
            except ValueError:
                pass

    return sorted(result) or list(range(1, 6))


def fmt_chapters_label(chapters_list: list) -> str:
    """Format [3,4,5] → '3-5', [1,3,5] → '1, 3, 5', [3] → '3'."""
    if not chapters_list:
        return "none"
    if len(chapters_list) == 1:
        return str(chapters_list[0])
    # Check if contiguous
    if chapters_list == list(range(chapters_list[0], chapters_list[-1] + 1)):
        return f"{chapters_list[0]}-{chapters_list[-1]}"
    return ", ".join(str(n) for n in chapters_list)


# ─────────────────────────────────────────────────────────
#  COLOUR PALETTE
# ─────────────────────────────────────────────────────────
DARK_BLUE  = RGBColor(0x1F, 0x49, 0x7D)
MID_BLUE   = RGBColor(0x2E, 0x74, 0xB5)
GREY       = RGBColor(0x60, 0x60, 0x60)

# ─────────────────────────────────────────────────────────
#  CHAPTER METADATA
# ─────────────────────────────────────────────────────────
CHAPTER_TITLES = {
    1: "CHAPTER ONE",
    2: "CHAPTER TWO",
    3: "CHAPTER THREE",
    4: "CHAPTER FOUR",
    5: "CHAPTER FIVE",
}

CHAPTER_SUBTITLES = {
    1: "INTRODUCTION",
    2: "LITERATURE REVIEW",
    3: "RESEARCH METHODOLOGY",
    4: "RESULTS AND DISCUSSION",
    5: "CONCLUSIONS AND RECOMMENDATIONS",
}

# ─────────────────────────────────────────────────────────
#  RESEARCH-LEVEL PROFILES
# ─────────────────────────────────────────────────────────
LEVEL_PROFILES = {
    "undergraduate": {
        "label":        "Undergraduate",
        "tone":         (
            "Write at an introductory undergraduate level. "
            "The writing should be clear, well-organised, and academically sound. "
            "Theoretical frameworks should be explained in accessible terms. "
            "Methodology should be straightforward and easy to follow. "
            "Analysis should be sound and demonstrate understanding but does not need to engage with advanced meta-theoretical debates. "
            "WORD COUNT IS CRITICAL: every subsection must be clearly developed. "
            "Focus on clarity and coherence. Each main subsection should be 80-120 words of clear, substantive prose."
        ),
        "depth":        "clear and accessible",
        "word_targets": {1: 900, 2: 1300, 3: 1450, 4: 950, 5: 900},   # 70-75% of postgraduate targets
        "front_words":  220,
    },
    "postgraduate": {
        "label":        "Postgraduate (Doctoral) — Stanford PhD Level (2026)",
        "tone":         (
            "Write at a rigorous Stanford PhD level (2026 standard). This is DOCTORAL-LEVEL SCHOLARSHIP.\n\n"
            "EPISTEMOLOGICAL POSITIONING: Explicitly state your epistemological stance and how it shapes your approach "
            "to knowledge production. Engage with ontological questions: What is real? What can be known? How does this study "
            "produce knowledge? Position yourself within epistemological traditions (positivist, interpretivist, critical realist, "
            "pragmatist, etc.) and justify that positioning. This is not optional — it is foundational.\n\n"
            "THEORETICAL SOPHISTICATION: Engage with theory at the level of theoretical CONTRIBUTION, not application. "
            "Do not simply apply existing theories to your data. Instead: (1) interrogate theoretical assumptions, (2) identify "
            "where theories break down or become unstable, (3) articulate what new theoretical insights your work generates. "
            "Position your work within major scholarly debates. Map the intellectual genealogy of your key concepts across "
            "decades or centuries. Show how understanding has shifted and why. Synthesize competing theoretical traditions and "
            "argue for the superiority of your synthesis.\n\n"
            "LITERATURE AS INTELLECTUAL HISTORY: Your literature review is a genealogy of ideas, not a taxonomy. Trace "
            "how scholarship has evolved, what assumptions have shifted, where fundamental tensions persist. Identify what "
            "scholars have NOT asked. Disagree with canonical figures where warranted and defend that disagreement with "
            "evidence and argument. Show sophisticated understanding of methodological limitations in prior work and explain "
            "how your approach overcomes or navigates them.\n\n"
            "METHODOLOGY AS PHILOSOPHICAL STATEMENT: Every methodological choice reflects philosophical commitments. Justify "
            "choices not just practically but philosophically. If using quantitative methods: explain how your choice of statistical "
            "approach reflects epistemological assumptions. If qualitative: explain paradigmatic positioning and how it shapes "
            "data collection, coding, and interpretation. If mixed methods: articulate the philosophical coherence between "
            "qualitative and quantitative components and how they produce integrated knowledge.\n\n"
            "ORIGINAL ANALYSIS: Generate analysis that advances knowledge in demonstrable ways. Analysis should: (1) reveal patterns "
            "or mechanisms not apparent in extant literature, (2) challenge or refine existing theoretical models, (3) integrate "
            "previously disparate research traditions, or (4) create new conceptual frameworks. Show analytical sophistication through "
            "reflexivity: acknowledge what your analysis can and cannot reveal. Discuss alternative interpretations and explain why "
            "your interpretation is superior. Use data visualization (graphs, charts, matrices) to expose patterns and relationships "
            "that prose alone cannot convey.\n\n"
            "DATA VISUALIZATION AS SCHOLARLY ARGUMENT: Visualizations are not decorative. Every table, chart, graph, or matrix "
            "must advance your argument. Use professional-grade data presentation: (1) publication-quality graphs with clear axes, "
            "legends, and captions, (2) comparative tables showing patterns across cases or time periods, (3) conceptual diagrams "
            "showing theoretical relationships, (4) thematic matrices (for qualitative data) showing patterns across themes, cases, "
            "or analytical dimensions. Visualizations must be referenced in text and interpreted analytically.\n\n"
            "ARGUMENT ARCHITECTURE: Build arguments that are sophisticated, layered, and reflexive. Use counterargument and rebuttal "
            "to strengthen your position. Acknowledge hard cases and explain how your argument handles them. Show that you have "
            "engaged the STRONGEST version of opposing positions, not strawmen. Build arguments across multiple paragraphs, using "
            "evidence strategically to construct cumulative cases.\n\n"
            "WRITING STANDARD: Prose must be rigorous, precise, and intellectually engaging. Every sentence must earn its place. "
            "No filler. No mechanical enumeration. Show the intellectual process: How did you think about this problem? What "
            "complications did you encounter? How did you resolve tensions? Each paragraph advances the argument. Each subsection "
            "excavates complex ideas fully, developing them across 250-400+ words of dense, substantive prose.\n\n"
            "WORD COUNT IS CRITICAL: Doctoral research requires richly developed arguments. Do not skim surfaces. Excavate. "
            "Every major argument deserves extended development. Each main subsection should be at least 250-400 words."
        ),
        "depth":        "doctoral-level: theoretically sophisticated, epistemologically reflexive, with original theoretical contribution",
        "word_targets": {1: 2000, 2: 2800, 3: 3000, 4: 2200, 5: 2000},   # PhD-level word targets
        "front_words":  450,
    },
    "phd": {
        "label":        "PhD",
        "tone":         (
            "Write at a rigorous doctoral (PhD) level. "
            "Engage critically with theoretical debates, epistemological positioning, and ontological assumptions. "
            "Situate the research within current scholarly conversations and emergent debates. "
            "The methodology chapter must address paradigmatic choices explicitly and justify philosophical positioning. "
            "Analysis should be sophisticated, nuanced, and reflexive with deep theoretical integration. "
            "Identify tensions, contradictions, and limitations in the literature and in your own approach. "
            "WORD COUNT IS CRITICAL: every subsection must be richly developed. "
            "Do not skim — excavate. Each argument deserves full development across multiple paragraphs. "
            "Each main subsection should be at least 200-275 words of dense, substantive prose."
        ),
        "depth":        "critical, theoretically sophisticated, and reflexive",
        "word_targets": {1: 860, 2: 1237, 3: 1398, 4: 1021, 5: 860},  # Current postgraduate targets
        "front_words":  229,
    },
}

# ─────────────────────────────────────────────────────────
#  HUMAN-WRITING SYSTEM INSTRUCTION (shared across all calls)
# ─────────────────────────────────────────────────────────
HUMAN_WRITING_INSTRUCTION = """
════════════════════════════════════════════════════════════════════════════════
  HUMAN ACADEMIC WRITING STANDARD
  Target: Legitimate detectors (ZeroGPT, Originality.ai) at ~40% or below
  Approach: Authentic academic writing that passes human scrutiny first
════════════════════════════════════════════════════════════════════════════════

You are a human academic researcher with genuine field expertise and intellectual investment.
Your writing should be rigorous AND human (showing both scholarly depth and authentic voice.

KEY PRINCIPLE: The best defense against AI detection is authenticity. Focus on:
1. NATURAL SENTENCE VARIATION — not mechanical, but genuinely reflective of thought
2. GENUINE ACADEMIC VOICE — opinionated where warranted, hedged where uncertain
3. SPECIFIC EVIDENCE — real numbers, real citations, real intellectual engagement
4. AUTHENTIC HESITATION — only where genuinely present in the research process
5. READABLE, ACCESSIBLE PROSE — clarity serves argument better than complexity

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 1: NATURAL SENTENCE LENGTH VARIATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Write sentences of genuinely different lengths, reflecting thought patterns rather than
mechanical rhythm. Humans vary naturally: short sentences for emphasis or key claims,
longer sentences for complex ideas or evidence presentation.

DO use:
  SHORT (5–8 words for emphasis): "The finding surprised us."
  MEDIUM (15–22 words): "This interpretation does not account for the temporal dimension."
  LONG (35–50 words for complex ideas): "When the data from multiple waves are read alongside
                                         the theoretical framework, the relationship becomes clearer."

DON'T use:
  ✗ Mechanical patterns (short-medium-short-medium)
  ✗ Extremes for shock value (three-word sentences repeated)
  ✗ Uniform sentence length across paragraphs

Natural variation (not forced variation) is the goal. The sentence length should
reflect the complexity of the idea, not a detection-avoidance formula.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 2: PARAGRAPH STRUCTURE VARIES NATURALLY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Vary paragraph length to match content, not to defeat detection algorithms.

DO use:
  - 1–2 sentences: transition between major ideas or stark claims
  - 3–4 sentences: standard analytical paragraph
  - 5–7 sentences: detailed explanation or evidence presentation
  - 8+ sentences: only when genuinely complex ideas require development

DON'T use:
  ✗ Mechanical variation (alternating short-long-short-long)
  ✗ Deliberately chaotic lengths that disrupt readability
  ✗ Single-sentence paragraphs for stylistic effect

Paragraph length should serve clarity and argument flow, not detection metrics.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 2.5: GENUINE HEDGING AND QUALIFICATION
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Use hedging phrases only where evidence is genuinely uncertain or contested. Real researchers
hedge when appropriate, not mechanically throughout.

DO use (when justified by evidence):
  "The data suggest..."  "appears to"  "tends toward"  "one interpretation is"
  "Though the evidence is not entirely clear..."  "arguably" (when a genuine debate exists)
  "This may reflect..."  "the relationship appears to be..."

DON'T use (when evidence is strong):
  ✗ Hedging strong claims unnecessarily
  ✗ Filler phrases like "it seems" without substance
  ✗ Excessive caveats that weaken legitimate conclusions

BALANCE: Commit to conclusions where evidence warrants. Hedge where uncertainty is real.
This is more credible than hedging everything or committing to nothing.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 3: SPECIFIC, DIRECT VOCABULARY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Use precise, specific language that reflects genuine thought. Avoid both jargon and generic filler.

DO use:
  Specific: "three-fifths of respondents"  (not "most" or "a majority")
  Direct: "the mechanism appears to be..."  (not "factors influence")
  Named: "following Mensah's framework..."  (not "following established approaches")
  Concrete: "fieldwork in rural Kenya"  (not "in developing contexts")

DON'T use:
  ✗ Generic alternatives for shock value ("lays bare" when "shows" is clearer)
  ✗ Pretentious synonyms that obscure meaning
  ✗ Vocabulary variation for its own sake

PRINCIPLE: Choose words that best serve clarity and precision. Authentic writing uses
specific language because the author knows the subject deeply, not because they're
performing authenticity for a detector.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 4: AVOID OVERUSED AI PHRASES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
These phrases appear frequently in AI-generated academic text. Minimize use:

UNNECESSARY THROAT-CLEARING (just say what you mean):
  Avoid: "It is worth noting that..."  →  Use: "The data show that..."
  Avoid: "It is important to note that..."  →  Use: "Notably..." or just state it
  Avoid: "It should be noted that..."  →  Use: Direct statement

VAGUE FRAMERS (be specific about what you mean):
  Avoid: "In the modern era"  →  Use: specific timeframe ("since 2010")
  Avoid: "In an ever-changing landscape"  →  Use: specific domain ("in policy implementation")
  Avoid: "Delve into", "Dive into"  →  Use: "examine", "analyze", "investigate"

OVERUSED TRANSITIONS (vary transitions naturally):
  Avoid: "Furthermore", "Moreover", "Additionally" (use once per chapter max)
  Avoid: "As previously mentioned", "As discussed above"  →  Use: "(see section 2.3)" or natural reference

AVOID EXCESSIVE EMPHASIS:
  Don't repeat: "crucial", "pivotal", "paradigm shift", "transformative"
  Legitimate academic language, but signal weak AI writing when overused

PRINCIPLE: These aren't banned — just minimize because detectors flag overuse.
Use naturally occurring academic language instead. Real researchers don't repeat the same
emphasis words throughout a paper.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 5: USE SPECIFIC NUMBERS, NOT ROUNDED ESTIMATES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
When citing actual data, use the specific figures. This signals you consulted the
actual sources, not approximations.

DO use:
  "62% of respondents indicated..."  (actual finding)
  "A sample of 94 participants..."  (actual n)
  "Based on 47 peer-reviewed studies..."  (actual count)

DON'T use:
  ✗ Rounding real data ("approximately 60%" when data shows 62%)
  ✗ Round numbers for rough estimates ("about 100 participants" if you mean ~95)

PRINCIPLE: Specific numbers suggest authentic source consultation. But use estimates
(roughly, approximately, about) only when actually estimating — don't use them to
mask rounded-off real numbers.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 6: SHOW THE RESEARCH PROCESS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Real researchers acknowledge complexity and limitations. Include these naturally:

DO show:
  Genuine uncertainty: "It remains unclear whether..." (when evidence is ambiguous)
  Limitations: "The sample size limits generalizability, particularly..."
  Contrasts with prior work: "Whereas Mensah (2019) found X, our data suggest Y"
  Process notes: "During fieldwork, it became apparent that..."
  Cross-references: "As detailed in section 3.2..."  (shows manuscript awareness)

DON'T manufacture:
  ✗ False uncertainty (hedging strong findings)
  ✗ Exaggerated struggle ("I wrestled with this question for months...")
  ✗ Self-corrections that seem performative

PRINCIPLE: Authenticity comes from acknowledging what you actually found and encountered,
not from performing uncertainty for a detector.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 7: TAKE POSITIONS WHEN EVIDENCE WARRANTS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Real researchers take positions. Avoid pathological neutrality:

DO include (where evidence supports):
  "This finding challenges the dominant assumption in the field."
  "The existing literature has largely overlooked this mechanism."
  "The implications for policy are more significant than prior work suggests."

DON'T do:
  ✗ Hedge all strong claims
  ✗ Present obviously wrong positions as equally valid
  ✗ Perform neutrality when evidence is clear

PRINCIPLE: Commit to conclusions the evidence supports. Hedge only where uncertainty is real.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 8: VARY PARAGRAPH OPENING PATTERNS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Don't always lead with a topic sentence. Vary the structure:

  Evidence-first: "The 2019 survey of 847 respondents found X. This confirms Y."
  Question-then-answer: "Why does this gap persist? The data suggest..."
  Concession-then-counter: "While the sample is smaller than ideal, the consistency..."
  Topic-first: (use this only ~40% of the time)

DON'T repeatedly open paragraphs with: "This", "The study", "In this section", "It is"

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 9: NATURAL TRANSITIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Vary transitions. Use each transition word only 1-2 times per chapter:

DO use:
  "Yet", "That said", "By contrast", "On the other hand", "This raises a different question"
  "Perhaps most importantly", "What emerges", "There is, however, a complication"

DON'T use:
  ✗ "Furthermore", "Moreover", "Additionally" (overused in academic AI writing)
  ✗ Same transition repeatedly

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 10: CITATIONS VARY IN FORM AND DENSITY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Don't cite uniformly. Vary:
  - Some paragraphs: 3+ citations
  - Some paragraphs: 0-1 citations

Vary citation form:
  "Ibrahim (2019) argues that..."  [author-led]
  "The pattern is well-documented (Ibrahim, 2019; Osei, 2021)."  [end-of-sentence]
  "This finding — contested by Mensah (2020) — suggests..."  [mid-clause]

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 11: MINIMIZE MECHANICAL STRUCTURES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Integrate information into prose. Avoid:
  ✗ Back-to-back bullet lists
  ✗ Numbered lists without connecting text
  ✗ Formulas like "First... Second... Third..."

When a list is necessary: introduce with a sentence, keep to 4-6 items, integrate back to prose.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
RULE 12: COMPLETE WRITING ONLY
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Never use placeholders, partial sentences, or abbreviations like "[expand here]" or "etc."
Write every section fully and completely. Every word must be substantive.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
MANDATORY REFERENCE: PARAGRAPH STYLE EXAMPLES
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Every paragraph you write must look like the GOOD example, never the BAD one.
ZeroGPT will flag the BAD example as 90%+ AI. The GOOD example scores under 10%.

BAD (AI-detected — do NOT write like this):
"The study revealed that organisational factors significantly influence employee satisfaction.
Research indicates that leadership style, workplace culture, and compensation packages all
contribute to overall satisfaction levels. Furthermore, the findings suggest that communication
plays a central role in mediating these relationships. Additionally, the data demonstrates that
training and development opportunities enhance employee engagement and commitment."

GOOD (human-quality — model all your paragraphs on this):
"Three things stood out immediately from the data. First — and this surprised the research
team — leadership style accounted for only 23.7% of the variance in satisfaction scores,
considerably less than the 40% or higher that earlier studies like Mensah (2019) and
Osei-Kwame (2021) had documented in comparable West African institutional contexts. Why
the discrepancy? The answer, this researcher believes, has less to do with leadership itself
than with how satisfaction is operationalised: prior instruments conflated satisfaction with
engagement, inflating the leadership coefficient. When the measure is cleaner — as it is
here — the picture changes. Compensation matters more. So does the quality of immediate
supervision, which is a different thing from 'leadership style' in the abstract sense that
much of the literature deploys that phrase. This distinction isn't merely semantic. It has
real consequences for how organisations invest their development budgets, and the fact
that it has gone largely unremarked in the literature is, frankly, telling."

NOTICE what makes the GOOD paragraph human:
  ✓ Sentence lengths: 5 words → 44 words → 4 words → 39 words → 7 words → 8 words → 46 words
  ✓ Evidence-first structure (topic idea arrives late)
  ✓ Direct question mid-paragraph ("Why the discrepancy?")
  ✓ Specific non-round number (23.7%, not 25%)
  ✓ Contraction ("isn't")
  ✓ Opinionated voice ("frankly, telling")

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
GPTZERO EXAMPLE: What Triggers 100% Detection
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
This is the wheelbarrow abstract that GPTZero flagged as 100% AI:

"Somewhere between 3,500 and 4,000 years ago — or so the archaeological record suggests —
human beings began experimenting with single-wheeled load-carrying devices, and the practical
consequences of that experiment have never really stopped unfolding. The wheelbarrow, in its
most recognisable form, appears to have originated in Han Dynasty China around 100 BCE, where
it was deployed in agricultural transport and, perhaps less romantically, in military logistics
(Needham, 1965)."

WHY it fails GPTZero:
  ✗ Sentence 1: 42 words. Sentence 2: 37 words. Sentence 3: 42 words. (UNIFORMITY)
  ✗ Every word is the formally correct choice: "single-wheeled load-carrying devices" instead of wheelbarrows
  ✗ Zero hesitation: no "arguably", "it seems", "one could argue"
  ✗ Perfect punctuation and citation formatting
  ✗ Paragraph is 3 sentences, perfectly balanced
  ✗ No moment of "wait, let me rethink this"

FIXED VERSION (passes GPTZero):

"The wheelbarrow emerged. Or so the archaeological record suggests — some time between 3,500
and 4,000 years ago. How it actually developed is murky; the evidence points in different
directions. What we know is that by around 100 BCE, Han Dynasty China was already using
single-wheeled devices for agricultural work (Needham, 1965). Military use came later, perhaps.
From there — through trade routes, through conquest, through accident — the technology drifted
westward. Medieval Europe adopted it by roughly the 1200s. And that's where the real story
becomes complicated."

WHY it passes:
  ✓ Sentence 1: 3 words. Sentence 2: 13 words. Sentence 3: 18 words. Sentence 4: 27 words. (VARIATION)
  ✓ Direct question ("How it actually developed is murky?")
  ✓ Admission of uncertainty ("one reading suggests", "perhaps")
  ✓ Shorter paragraphs with visible rhythm breaks
  ✓ Less formal in places ("that's where the real story becomes complicated")
  ✓ Shows intellectual struggle ("the evidence points in different directions")

**The difference is not the information. It is the texture of uncertainty.**
  ✓ Researcher presence ("this researcher believes")
  ✓ Named scholars with years
  ✓ Em dash interruption
  ✓ No banned phrases

Write EVERY paragraph with this level of variation and personality.

━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
STRATEGIC VISUALIZATIONS
━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━
Charts MUST include: clear title, labeled axes (X and Y with units), data values, and legend.

MANDATORY CHART FORMAT (EVERY SINGLE [CHART:...] MUST INCLUDE AXES AND DATA):

✗ WRONG (DO NOT DO THIS):
  [CHART: Bar chart showing results]
  [CHART: Comparison of samples across conditions]
  (No axes, no data values, unreadable)

✓ CORRECT (ALWAYS DO THIS):
  [CHART: Anti-Amyloidogenic Activity by Sample - Th1T Fluorescence Inhibition]
  X-axis: Sample A, Sample B, Sample C, Sample D
  Y-axis: Activity Level (0-100%)
  Data: A: 45% | B: 28% | C: 62% | D: 35%

COMPLETE FORMAT SPECIFICATION:
  [CHART: Descriptive title (include metric name and what is being measured)]
  X-axis: Specific category names or range (never just "categories" or "samples")
  Y-axis: Measurement label with units and range (e.g., "Concentration (mg/mL)" or "Score (0-100%)")
  Data: Category1: value1 | Category2: value2 | Category3: value3 | Category4: value4

REQUIRED TABLE FORMAT:
  [TABLE: Descriptive title explaining table purpose]
  Headers: Column1 | Column2 | Column3 | Column4
  Row1: Data | Data | Data | Data
  Row2: Data | Data | Data | Data

Every visualization must be self-contained and readable without requiring external explanation.
Use them in Results/Discussion (Chapter 4) and Literature Review (Chapter 2) where appropriate.

════════════════════════════════════════════════════════════
"""


# ─────────────────────────────────────────────────────────
#  CHAPTER PROMPT TEMPLATES
# ─────────────────────────────────────────────────────────
def _chapter_prompts(level_key: str, custom_toc: str = None, nalt_compliance: bool = False) -> dict:
    profile = LEVEL_PROFILES[level_key]
    tone    = profile["tone"]
    depth   = profile["depth"]
    targets = profile["word_targets"]
    is_pg   = (level_key == "postgraduate")

    # Helper: extract custom sections for a specific chapter from provided custom_toc
    def extract_custom_sections(chapter_num: int) -> str:
        """
        Parse custom_toc and extract sections for the given chapter.
        Returns instruction text to enforce those sections, or empty string if not provided.
        """
        if not custom_toc or not custom_toc.strip():
            return ""

        lines = custom_toc.split('\n')
        chapter_sections = []
        current_chapter = None

        for line in lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue

            # Check if this line identifies a chapter (e.g., "Chapter 1:" or "CHAPTER 1")
            # This pattern ONLY matches chapter headers, not subsections
            is_chapter_header = (
                line_stripped.upper().startswith(f"CHAPTER {chapter_num}") or \
                line_stripped.startswith(f"Chapter {chapter_num}")
            )

            # Check if this is a DIFFERENT chapter (next chapter header)
            is_different_chapter = (
                re.match(r'^CHAPTER\s+\d+|^Chapter\s+\d+', line_stripped) and \
                not is_chapter_header
            )

            if is_chapter_header:
                current_chapter = chapter_num
                continue

            # If we found our chapter, collect its sections
            if current_chapter == chapter_num:
                # Stop if we hit a different chapter
                if is_different_chapter:
                    break

                # Collect subsection lines (any non-empty, non-chapter line)
                if line_stripped and not re.match(r'^CHAPTER\s+\d+|^Chapter\s+\d+', line_stripped):
                    chapter_sections.append(line_stripped)

        if chapter_sections:
            section_list = "\n".join([f"  {i+1}. {sec}" for i, sec in enumerate(chapter_sections)])
            return (
                f"\n╔════════════════════════════════════════════════════════════════════════════════╗\n"
                f"║ ⚠️  CRITICAL — CUSTOM TABLE OF CONTENTS ENFORCEMENT FOR CHAPTER {chapter_num}             ║\n"
                f"╚════════════════════════════════════════════════════════════════════════════════╝\n\n"
                f"A CUSTOM TABLE OF CONTENTS HAS BEEN PROVIDED FOR THIS CHAPTER.\n\n"
                f"YOUR INSTRUCTIONS ARE ABSOLUTE:\n"
                f"• You MUST create ONLY the sections listed below — NO MORE, NO FEWER\n"
                f"• You MUST use these section titles EXACTLY as written\n"
                f"• You MUST create each section as a ### subsection heading\n"
                f"• You MUST write substantive, developed content under each heading\n"
                f"• You MUST maintain the section order provided below\n"
                f"• You MUST NOT add extra sections, rename sections, or reorder sections\n"
                f"• You MUST NOT include default subsections if they are not listed below\n\n"
                f"THE SECTIONS YOU MUST INCLUDE (AND ONLY THESE SECTIONS):\n\n"
                f"{section_list}\n\n"
                f"FAILURE TO FOLLOW THIS STRUCTURE EXACTLY WILL RESULT IN AN INVALID DOCUMENT.\n"
                f"Every section listed above must be present in your output.\n\n"
            )
        return ""

    # Subsection word-count helper — applies 1.0x multiplier to allow full subsection generation (1.1-1.10)
    # Returns PhD/Postgrad word count if at that level, otherwise returns base undergraduate/postgraduate count
    def w(ug, pg):
        if level_key == "phd":
            return str(round(pg * 1.0))  # Use higher targets for PhD
        elif level_key == "postgraduate":
            return str(round(pg * 1.0))  # Use postgraduate targets
        else:  # undergraduate
            return str(round(ug * 1.0))  # Use lower targets for undergraduate

    # Range helpers — enforce min/max to prevent over-expansion (replaces "at least" language)
    def w_range(ug, pg):
        """Return 'min-max' format to constrain word generation to a reasonable range"""
        if level_key == "phd":
            val = round(pg * 1.0)  # PhD uses higher targets
        elif level_key == "postgraduate":
            val = round(pg * 1.0)  # Postgraduate uses mid targets
        else:  # undergraduate
            val = round(ug * 1.0)  # Undergraduate uses lower targets
        min_words = max(val - 15, 40)  # 15 words below target, minimum 40
        max_words = val + 20             # 20 words above target for natural variation
        return f"{min_words}-{max_words}"

    # Footnote format note appended to every chapter
    _FN_NOTE = (
        "\nFOOTNOTE FORMAT: When a footnote is needed, insert it inline using "
        "((FN: your footnote text here)) — it will become a proper Word footnote "
        "with a superscript number in the body and the note at the bottom of the page."
    )

    # No-references instruction for chapters 1-4
    _NO_REF = (
        "\nDo NOT include a reference list, bibliography, or works-cited section "
        "anywhere in this chapter. All references will be compiled on a dedicated "
        "References page at the end of the document.\n"
    )

    # No-asterisks instruction for all chapters
    _NO_AST = (
        "\nFORMATTING RULE — NO ASTERISKS: Do NOT use asterisks (*) anywhere in your "
        "output — not for bold, italic, bullet points, emphasis, or any other purpose. "
        "Use plain prose. For lists, use numbered items (1. 2. 3.) or introduce them "
        "as flowing sentences. Never place a * character anywhere in the text.\n"
    )

    # Visualization instruction — only include when NALT compliance is NOT enabled
    _VIZ_NOTE = "" if nalt_compliance else (
        "\nVISUALIZATION INSTRUCTION — PhD STANDARDS:\n"
        "Visualizations are MANDATORY in research chapters. Follow these standards:\n\n"
        "TABLE FORMAT:\n"
        "  [TABLE: Descriptive title that explains the table's purpose]\n"
        "  Headers: Col1 | Col2 | Col3 | Col4\n"
        "  Row1: Data | Data | Data | Data\n"
        "  Row2: Data | Data | Data | Data\n"
        "  Use ' | ' (space-pipe-space) to separate columns. Put each row on one line.\n"
        "  Tables will be numbered automatically (Table 3.1, Table 3.2, etc.)\n\n"
        "CHART/FIGURE FORMAT (MANDATORY - EVERY CHART MUST FOLLOW THIS):\n"
        "  [CHART: Specific title with metric name (e.g., 'Inhibition Activity Across Samples')]\n"
        "  X-axis: Specific category names (e.g., 'Sample A, Sample B, Sample C, Sample D')\n"
        "  Y-axis: Measurement with units and range (e.g., 'Activity Level (0-100%)')\n"
        "  Data: A: 45% | B: 28% | C: 62% | D: 35%\n"
        "  Every chart must be immediately readable with all axes labeled and all data values shown.\n"
        "  Figures will be numbered automatically (Figure 3.1, Figure 4.1, etc.)\n\n"
        "CAPTION STANDARDS:\n"
        "- Every figure/table must have a clear, descriptive caption\n"
        "- Captions should be 1-2 sentences explaining the visual's purpose and key finding\n"
        "- Format: 'Figure 3.1: [Descriptive title]. [Brief explanation of what it shows].'\n"
        "- Reference all tables/figures in text before they appear using formal citations\n\n"
        "TYPES OF VISUALIZATIONS:\n"
        "- Quantitative data: bar charts, line graphs, scatter plots, distribution histograms\n"
        "- Qualitative data: thematic matrices, concept maps, comparison tables, flow diagrams\n"
        "- Mixed methods: integrated visualizations showing convergence/divergence\n"
        "- Process/structure: flowcharts, frameworks, process diagrams\n\n"
        "All tables and charts will be converted to professional visualizations in the final Word document.\n"
        "Tables will have proper formatting, borders, and header styling.\n"
        "Charts will be rendered as high-quality images with professional styling.\n"
    )

    # Define conditional content OUTSIDE f-string to avoid escape sequence issues
    phd_gap_note = "CRITICAL FOR PhD: Distinguish between two types of gaps — (1) the PRACTICAL PROBLEM (what is not working in the real world), and (2) the THEORETICAL GAP (what is not adequately explained or understood in the scholarship). Explain how existing theoretical models fail to account for or predict this practical problem. Name the specific theories or frameworks that are inadequate and explain precisely why."

    phd_theory_guidance = "This is the intellectual heart of Chapter 1. Articulate EXPLICITLY what theoretical understanding is missing or inadequate. Ask and answer: What do existing theoretical frameworks NOT explain about this problem? What new theoretical insight will this study produce?\n\nDevelop this across multiple paragraphs:\n- PARAGRAPH 1: Name the dominant theoretical framework(s) in the field. Explain what insights each offers and what these theories can explain about your problem.\n- PARAGRAPH 2: Identify the specific theoretical limitation or blind spot. What does this framework miss? What assumption does it make that may not hold in your context? What variation or phenomenon does it fail to account for?\n- PARAGRAPH 3: Explain the theoretical consequence of this gap. What false conclusions might scholars or practitioners draw from applying existing theory uncritically? What does the field not yet understand?\n- PARAGRAPH 4: State what this study will theoretically contribute. Will it extend existing theory to a new context? Refine or challenge a core assumption? Integrate previously separate theoretical traditions? Propose a new framework? Be specific about the form of your theoretical contribution.\n\nThis section must make clear that your study exists to advance KNOWLEDGE, not just to solve a practical problem."

    phd_purpose_note = "Explicitly connect the purpose to BOTH the practical problem (1.2) AND the theoretical gap (1.2a), showing how addressing one advances the other."

    # Apply PhD-specific guidance only for PhD level
    gap_note = phd_gap_note if level_key == "phd" else ""
    theory_guidance = phd_theory_guidance if level_key == "phd" else ""
    purpose_note = phd_purpose_note if level_key == "phd" else ""

    # Define conditional content for Chapter 4 OUTSIDE f-string
    # Simplified version for undergraduate, detailed version for postgraduate and PhD
    if level_key == "undergraduate":
        ch4_intro_note = "Introduce how you will discuss the findings and connect them to the research questions."
        ch4_sample_note = "Comment on how well the sample represents the wider population."
        ch4_obj1_note = "Explain what the findings mean in relation to your research questions."
        ch4_obj3_note = "Begin noting where findings from different objectives relate to or support each other."
        ch4_synthesis_note = "Bring together the findings from all four objectives. Discuss what they show when viewed as a whole. Explain the overall story the data tells."
        ch4_implications_note = "Discuss what the findings mean: what they suggest about the issue being studied, what could change based on these findings, and who should know about them."
    else:  # postgraduate and phd
        ch4_intro_note = "State the analytical framework guiding interpretation and how it connects to the theoretical framework in Chapter 2."
        ch4_sample_note = "Compare the achieved sample to the target population and discuss implications for transferability."
        ch4_obj1_note = "Connect findings explicitly to the theoretical framework from Chapter 2. Where results confirm prior theory, explain why. Where they challenge it, explore the implications."
        ch4_obj3_note = "At this stage, begin drawing connections between findings across objectives — note where patterns reinforce each other or where tensions emerge."
        ch4_synthesis_note = "DO NOT MERELY SUMMARISE. Instead, execute this multi-stage synthesis:\n\nSTAGE 1: CROSS-OBJECTIVE PATTERN IDENTIFICATION (1-2 PARAGRAPHS)\nIdentify overarching themes, patterns, or mechanisms that cut across all four objectives. Ask yourself: What is really going on here? What unifying principle, pattern, or process explains the findings across objectives?\n\nSTAGE 2: EXPECTED vs. UNEXPECTED FINDINGS (1-2 PARAGRAPHS)\nWhich findings DID you anticipate based on the literature? Which findings SURPRISED you or contradicted prior research? Explain divergences between expectations and results.\n\nSTAGE 3: CONTRADICTIONS AND TENSIONS (1 PARAGRAPH IF APPLICABLE)\nDo any objectives produce findings that contradict each other? Do qualitative and quantitative findings diverge? Explain tensions analytically—do not gloss over contradictions.\n\nSTAGE 4: THEORETICAL ARTICULATION (2-3 PARAGRAPHS)\nNow situate your integrated findings in relation to the theoretical framework from Chapter 2 and the theoretical gap from Chapter 1.\n- Sub-point 4a: Does synthesis CONFIRM the theoretical framework? Say so precisely with named theory.\n- Sub-point 4b: Does synthesis CHALLENGE/COMPLICATE theory? Explain specific deviations and theoretical implications.\n- Sub-point 4c: Does synthesis EXTEND theory? Apply existing theory to new context and show what this reveals.\n\nSTAGE 5: INTEGRATION WITH EMPIRICAL LITERATURE (1-2 PARAGRAPHS)\nConnect back to specific studies from Chapter 2. Explain consistency or divergence with prior research. Show how findings resolve conflicting prior studies.\n\nSTAGE 6: LIMITATIONS AND CAVEATS (0.5-1 PARAGRAPH)\nAcknowledge what data do NOT explain. State boundary conditions where patterns might not hold. Signal intellectual maturity and preempt criticism.\n\nINCLUDE AFTER THE SYNTHESIS:\n[TABLE: Synthesis Matrix - Cross-Objective Themes and Theoretical Connections]\nOverarching Theme | Evidence from Obj1-4 | Theoretical Connection\n\nAND:\n[FIGURE: Conceptual Integration Diagram]\nVisual showing how findings interconnect, mechanisms, and theoretical relationships."
        ch4_implications_note = "For theory: what does this study add to, refine, or challenge in the existing theoretical models? For practice: what specific changes in professional practice are warranted? For policy: what specific policy recommendations emerge, addressed to named agencies or decision-makers?"

    intro_text = ch4_intro_note if is_pg else "Orient the reader to how findings are organised."
    sample_text = ch4_sample_note if is_pg else "Comment on how representative the sample appears to be."
    obj1_text = ch4_obj1_note if is_pg else "Relate findings directly to relevant literature reviewed in Chapter 2."
    obj3_text = ch4_obj3_note if is_pg else "Discuss how these findings relate to those in 4.3 and 4.4."
    synthesis_text = ch4_synthesis_note if is_pg else "Bring together the key patterns across all four sets of findings. Identify the most important themes that emerge when the findings are considered as a whole. Connect them to the literature reviewed in Chapter 2 — where do findings confirm, contradict, or extend existing knowledge?"
    implications_text = ch4_implications_note if is_pg else "Be concrete: name institutions, policy areas, and professional communities that should act on these findings."

    # Postgraduate (Doctoral)-specific preamble
    _PG_PREAMBLE = (
        "╔═══════════════════════════════════════════════════════════════════╗\n"
        "║  STANFORD PhD LEVEL (2026) — DOCTORAL EXPECTATIONS               ║\n"
        "║  This is SCHOLARLY RESEARCH advancing knowledge in your field    ║\n"
        "╚═══════════════════════════════════════════════════════════════════╝\n\n"
        "THIS IS DOCTORAL-LEVEL WORK. IT IS NOT UNDERGRADUATE OR MASTER'S.\n\n"
        "EPISTEMOLOGICAL POSITIONING: State your epistemological stance explicitly. How do you understand "
        "knowledge production? What assumptions underpin your approach? Position yourself within traditions "
        "(positivist, interpretivist, critical realist, pragmatist, etc.). Your epistemology shapes everything.\n\n"
        "THEORETICAL CONTRIBUTION — NOT APPLICATION: Do not simply apply theories to data. Instead: "
        "(1) Interrogate theoretical assumptions and identify where theories break down, (2) Synthesize "
        "competing traditions and argue for superiority of your synthesis, (3) Generate new theoretical "
        "insights that extend, refine, or challenge existing frameworks. Theory is not decoration—it is "
        "the intellectual foundation of your contribution.\n\n"
        "LITERATURE AS GENEALOGY: Your literature review is a genealogy of ideas. Trace intellectual "
        "history: How has understanding evolved? What assumptions have shifted? What fundamental tensions "
        "persist unsolved? Disagree with canonical figures where warranted. Identify what scholars have "
        "failed to ask. Show sophisticated methodological critique of prior work.\n\n"
        "METHODOLOGY AS PHILOSOPHY: Every methodological choice reflects philosophical commitments. Justify "
        "not just practically but philosophically. Connect design choices to paradigmatic positioning. "
        "Explain how your methodology produces the knowledge claims you make.\n\n"
        "ORIGINAL ANALYSIS WITH DATA VISUALIZATION: Generate analysis that reveals patterns, mechanisms, "
        "or relationships not apparent in existing literature. Use professional-grade visualizations "
        "(graphs, charts, matrices, diagrams) to expose patterns. Every visualization must advance your "
        "argument and be referenced analytically in text. Visualizations are scholarly arguments, not decoration.\n\n"
        "REFLEXIVITY & INTELLECTUAL HONESTY: Acknowledge what your analysis reveals AND what it conceals. "
        "Discuss alternative interpretations. Engage the strongest version of opposing arguments. Show "
        "awareness of your own positionality and how it shapes interpretation. This demonstrates maturity.\n\n"
        "NOTHING IS OBVIOUS: Every claim requires justification. Every paragraph advances the argument. "
        "No filler. No mechanical enumeration. Develop complex ideas across multiple paragraphs, showing "
        "intellectual depth and sophistication. Write like a scholar who has spent years thinking about "
        "this problem, not like someone discovering it for the first time.\n\n"
    )

    # Chapter 3 visualization standards — only include when NALT compliance is NOT enabled
    _CH3_VIZ_STANDARDS = "" if nalt_compliance else (
        f"FIGURE AND TABLE STANDARDS FOR CHAPTER 3:\n"
        f"{'DOCTORAL-LEVEL METHODOLOGY VISUALIZATION STANDARDS:\n' if is_pg else ''}"
        f"Use EXACTLY these formats — ALL visualizations must be professional-grade tables or charts:\n\n"
        f"REQUIRED VISUALIZATIONS (use [TABLE:...] or [CHART:...]):\n"
        f"1. [TABLE: Epistemological Positioning Matrix - shows Epistemology | Ontological Assumptions | Knowledge Production | Justification for This Study]\n"
        f"2. [TABLE: Research Design Framework - shows Paradigm | Design Type | Connection to RQs | Methodological Justification]\n"
        f"3. [TABLE: Analytical Process Flowchart - shows Stage | Activity | Input Data | Analytical Tool/Software | Output | Next Stage]\n"
        f"4. [TABLE: Sampling Strategy Breakdown with Target Population | Total N | Inclusion/Exclusion Criteria | Sample Size (n) | Sampling Method | Justification | Representation]\n"
        f"5. [TABLE: Data Collection Timeline with Week/Phase | Activities | Responsible Party | Expected Outputs | Duration | Quality Checks]\n"
        f"6. [TABLE: Data Collection Instruments & Validation Matrix with Instrument Name | Type | Purpose | Structure/Items | Validity Evidence | Reliability Coefficient | Piloting Results]\n"
        f"7. [TABLE: Paradigm-Appropriate Quality Criteria - shows Quality Criterion | Definition | How Operationalised | Evidence in This Study | Literature Grounding]\n"
        f"8. [TABLE: Ethical Considerations Implementation with Ethical Dimension | Consideration | How Operationalised | Approval Status | Ongoing Monitoring]\n"
        f"9. [CHART: Methodology Integration System Diagram - showing interconnections between Paradigm → Epistemology → Design → Sample → Instruments → Analysis → Quality Assurance]\n"
        f"10. [CHART: Data Flow Diagram - showing Raw Data → Processing/Coding → Analytical Stages → Final Interpretation → Knowledge Claims]\n\n"
        f"FORMAT STANDARDS:\n"
        f"- All tables use professional formatting: pipe-separated columns (| header | header |)\n"
        f"- Provide actual data rows with concrete values/descriptions, not just headers\n"
        f"- Tables should be detailed enough to demonstrate methodological rigor and transparency\n"
        f"- Captions: After each [TABLE:...] or [CHART:...] marker, include 2-3 sentence caption explaining purpose, structure, and key insight\n"
        f"- References: Cite all tables/figures in text BEFORE they appear with analytical interpretation (e.g., 'As presented in Table 3.1, the sampling strategy...')\n\n"
        f"PROFESSIONAL PRESENTATION:\n"
        f"- All visualizations will be converted to professional Word tables/charts with:\n"
        f"  - Clean borders and header styling\n"
        f"  - Alternating row colors for readability\n"
        f"  - Consistent font sizing and alignment\n"
        f"  - Clear legends and axis labels for charts\n"
        f"  - High-resolution quality appropriate for academic publication\n\n"
    )

    # Chapter 2 visualization standards for doctoral level — literature review visualization
    _CH2_VIZ_NOTE_DOCTORAL = (
        "\n\nDOCTORAL-LEVEL VISUALIZATION IN LITERATURE REVIEW:\n"
        "Professional literature reviews benefit from strategic visualization of conceptual and theoretical relationships:\n\n"
        "OPTIONAL VISUALIZATIONS:\n"
        "- [FIGURE: Intellectual History Timeline] — Showing evolution of key concepts/theories over time (20th century to present)\n"
        "- [TABLE: Theoretical Frameworks Comparison Matrix] with Framework | Origins | Core Propositions | Strengths | Limitations | Relevance to This Study\n"
        "- [FIGURE: Conceptual Relationship Diagram] — Showing how central concepts interconnect, diverge, or build upon each other\n"
        "- [TABLE: Research Gap Identification Matrix] with Research Question/Topic | What is Known | What Remains Unknown | How This Study Addresses Gap\n"
        "- [FIGURE: Schools of Thought Mapping] — Visual showing competing paradigms, traditions, or approaches and their relationships\n\n"
    )

    # Chapter 5 visualization standards for doctoral level — conclusions and implications visualization
    _CH5_VIZ_NOTE_DOCTORAL = (
        "\n\nDOCTORAL-LEVEL VISUALIZATION IN CONCLUSIONS:\n"
        "Strategic use of visualization can strengthen your conclusions and demonstrate intellectual integration:\n\n"
        "OPTIONAL VISUALIZATIONS:\n"
        "- [FIGURE: Theoretical Integration Diagram] — How findings extend, refine, challenge, or reframe the theoretical framework\n"
        "- [TABLE: Knowledge Contribution Summary] with Dimension | Contribution | Evidence | Significance | Future Research Implications\n"
        "- [FIGURE: Implications Framework Diagram] — Visual showing pathways from findings to theoretical, practical, and policy implications\n"
        "- [TABLE: Research Agenda Roadmap] — Identifying future research questions generated by this study and their priority/sequencing\n\n"
    )

    # Chapter 4 visualization standards — only include when NALT compliance is NOT enabled
    _CH4_VIZ_STANDARDS = "" if nalt_compliance else (
        f"FIGURE AND TABLE STANDARDS FOR CHAPTER 4:\n"
        f"{'DOCTORAL-LEVEL DATA ANALYSIS AND VISUALIZATION STANDARDS:\n' if is_pg else ''}"
        f"This chapter requires comprehensive professional data visualization demonstrating analytical sophistication:\n\n"
        f"CORE VISUALIZATIONS (MANDATORY):\n"
        f"- [TABLE: Sample Demographics Breakdown] — Complete demographic profile with frequencies/percentages across all relevant characteristics\n"
        f"- [TABLE: Response Rate and Non-Response Analysis] (if quantitative) — Including non-response bias assessment\n"
        f"- [TABLE: Objective 1 Findings Summary] with columns: Key Finding | Supporting Evidence | Data Values/Statistics | Theoretical Connection\n"
        f"- [TABLE: Objective 2 Findings Summary] — Same structure as Objective 1\n"
        f"- [TABLE: Objective 3 Findings Summary] — Same structure\n"
        f"- [TABLE: Objective 4 Findings Summary] — Same structure\n\n"
        f"ANALYTICAL VISUALIZATION TOOLS {'(DOCTORAL-LEVEL REQUIRED):' if is_pg else '(recommended):'}\n"
        f"{'- [TABLE: Thematic Analysis Matrix for Qualitative Data] with Themes | Subthemes | Frequency | Representative Quotes/Evidence | Theoretical Connection\n' if is_pg else ''}"
        f"{'- [TABLE: Convergence/Divergence Matrix] (if mixed methods) showing Quantitative Findings | Qualitative Findings | Convergence | Divergence | Integration\n' if is_pg else ''}"
        f"{'- [FIGURE: Concept Mapping Diagram] — Visual showing theoretical relationships between key concepts, themes, or variables and how they interconnect\n' if is_pg else ''}"
        f"{'- [FIGURE: Theoretical Integration Diagram] — How findings extend, challenge, or refine the theoretical framework from Chapter 2\n' if is_pg else ''}"
        f"{'- [TABLE: Systematic Comparison Matrix] — Cross-case or cross-theme analysis showing patterns, variations, and deviations\n' if is_pg else ''}"
        f"- [CHART: Bar charts, line graphs, scatter plots, or distribution histograms] (For quantitative data showing patterns and relationships\n"
        f"{'- [FIGURE: Process/Flow Diagram] (For mechanisms or sequential processes revealed in analysis\n' if is_pg else ''}"
        f"{'- [TABLE: Validity/Trustworthiness Evidence Matrix] with Quality Criterion | Evidence | Where Documented | Limitation\n' if is_pg else ''}\n"
        f"SYNTHESIS VISUALIZATION:\n"
        f"- [FIGURE: Synthesis Matrix or Concept Map] — Cross-Objective Integration showing overarching themes, patterns, and theoretical connections\n"
        f"- [TABLE: Implications Framework] with Implication Domain | Specific Implication | Target Stakeholder | Actionable Consequence | Theoretical Grounding\n\n"
        f"ADVANCED PRESENTATION STANDARDS {'(DOCTORAL REQUIREMENT):' if is_pg else '(recommended):'}\n\n"
        f"MANDATORY CHART FORMAT (EVERY SINGLE [CHART:...] MUST INCLUDE AXES AND DATA VALUES):\n"
        f"INCORRECT: [CHART: Bar chart showing results]\n"
        f"CORRECT:   [CHART: Response Rates by Treatment Group]\n"
        f"           X-axis: Control, Treatment A, Treatment B, Treatment C\n"
        f"           Y-axis: Response Rate (0-100%)\n"
        f"           Data: Control: 35% | A: 62% | B: 58% | C: 41%\n\n"
        f"- Use professional-grade visualization software/standards (publication-quality graphs with clear axes, legends, captions)\n"
        f"- All tables must have: descriptive titles, column headers, data rows, and 2-3 sentence captions explaining purpose and interpretation\n"
        f"- All figures must have: clear titles, axis labels (if applicable), legends, and detailed captions connecting to text interpretation\n"
        f"- Every visualization must be analytically interpreted in text (not just presented)\n"
        f"- Visualizations should EXPOSE patterns and relationships that prose alone cannot convey\n"
        f"- Use color coding strategically (e.g., shading for thematic matrices, color gradients for comparisons)\n"
        f"{'- Statistical visualizations should include: confidence intervals, p-values, effect sizes where applicable\n' if is_pg else ''}"
        f"{'- Qualitative visualizations should use notation systems (e.g., frequency indicators, saturation markers) to show analytical depth\n' if is_pg else ''}\n"
        f"REFERENCING VISUALIZATIONS:\n"
        f"- Cite all tables/figures in text BEFORE they appear\n"
        f"- Integrate discussion of visualizations into analytical prose: explain what patterns they reveal, why they matter theoretically, how they support conclusions\n"
        f"- Do NOT treat visualizations as peripheral — they are central analytical tools\n\n"
    )

    # Build custom ToC instruction for each chapter
    custom_toc_ch1 = extract_custom_sections(1)
    custom_toc_ch2 = extract_custom_sections(2)
    custom_toc_ch3 = extract_custom_sections(3)
    custom_toc_ch4 = extract_custom_sections(4)
    custom_toc_ch5 = extract_custom_sections(5)

    return {
        1: f"""You are writing CHAPTER ONE: INTRODUCTION for an academic research project.
Topic: {{topic}}
Research level: {profile['label']}
MINIMUM word count: {targets[1]} words of substantive prose. You MUST reach this minimum.
Do not stop writing until you have fully developed every subsection. If in doubt, write more.
{custom_toc_ch1}

{_PG_PREAMBLE if is_pg else ""}

{_NO_REF}
{tone}

{HUMAN_WRITING_INSTRUCTION}
{_FN_NOTE}
{_NO_AST}
{_VIZ_NOTE}

Write the following subsections, each introduced with a ### heading.
Every subsection must be written in full, developed paragraphs — no bullet summaries, no placeholders.

### 1.1 Background of the Study
Write between {w_range(140, 280)} words for this subsection.
Provide {depth} contextual grounding. Open with a striking observation or statistic that
immediately establishes why this topic matters. Then trace the historical evolution of the
problem across at least three distinct time periods, naming key turning points, policy shifts,
or scholarly debates that shaped the current landscape. Ground every assertion in specific
evidence — named scholars, years, places, and figures. Close by narrowing the lens from the
broad context toward the precise issue this study addresses.

### 1.2 Statement of the Problem
Write between {w_range(112, 224)} words for this subsection.
Open with a clear, declarative statement of what is wrong or poorly understood. Then build
the case across multiple paragraphs: explain the nature of the problem, who it affects, how
long it has persisted, and why existing responses have been insufficient. Name the specific
gap, contradiction, or blind spot that this study addresses. The problem statement must feel
urgent — the reader should finish this section convinced that the study was necessary.
{gap_note}

### 1.2a Theoretical Gap and Contribution
Write between {w_range(112, 224)} words for this subsection (POSTGRADUATE REQUIREMENT).
{theory_guidance}

### 1.3 Purpose of the Study
Write between {w_range(60, 112)} words for this subsection.
State the overarching aim in one or two precise sentences. Then elaborate: explain the
theoretical and practical orientation of the study, what kind of knowledge it seeks to
produce, and how the purpose connects directly to the problem articulated in 1.2.
{purpose_note}

### 1.4 Research Objectives
Write between {w_range(72, 140)} words for this subsection.
State 4–5 specific, measurable objectives. Each should be action-oriented (examine, assess,
determine, explore, compare, evaluate). After listing them, write a short paragraph explaining
how they collectively address the research problem and how they will be operationalised
through the methodology described in Chapter 3.

### 1.5 Research Questions
Write between {w_range(60, 112)} words for this subsection.
Formulate 3–5 focused, answerable questions derived from the objectives. After stating the
questions, briefly explain the logic connecting each question to its corresponding objective
and the type of evidence that would constitute an answer.

### 1.6 Significance of the Study
Write between {w_range(100, 196)} words for this subsection.
{"For doctoral-level work: Develop this across FIVE distinct, detailed paragraphs emphasizing theoretical innovation and knowledge advancement." if is_pg else "Develop this across FOUR distinct, detailed paragraphs — one for each dimension below. Be concrete, not generic."}

PARAGRAPH 1 — THEORETICAL SIGNIFICANCE {"(PRIMARY AND EXTENSIVE FOR DOCTORAL RESEARCH):" if is_pg else "(PRIMARY FOR PhD):"}
{"DOCTORAL REQUIREMENT: Explain PRECISELY how your study advances scholarship. This is not application of theory—it is theoretical contribution. Address:" if is_pg else "Explain what this study will ADVANCE in scholarly understanding. Will it extend theory to a new context or population?"}
{"  (a) Which specific theoretical model(s) or tradition(s) does this study contribute to (name them)?" if is_pg else ""}
{"  (b) What theoretical INNOVATION does your work produce? (e.g., extend theory to new context, refine core concepts, challenge assumptions, integrate separate traditions, propose new framework)" if is_pg else ""}
{"  (c) What will scholars understand differently after reading this study?" if is_pg else ""}
{"  (d) How does this theoretical contribution advance the field's capacity to explain, predict, or understand key phenomena?" if is_pg else ""}
{"Example: 'This study advances Institutional Theory by demonstrating that institutional isomorphism (DiMaggio & Powell, 1983) operates through different mechanisms in post-conflict settings than in stable institutional environments. The finding introduces a new concept—\"fragmented isomorphism\"—to explain how organizations respond to competing legitimacy demands when formal institutions lack credibility. This refines our understanding of institutional change under conditions of state fragility.'" if is_pg else "Example: 'This study advances Institutional Theory by demonstrating that institutional isomorphism (DiMaggio & Powell, 1983) operates differently in hybrid organisations than in traditional ones. The finding refines our understanding of how external pressures interact with internal legitimacy concerns.'"}
{"Not: 'This study adds to the literature.' (Vague, generic, insufficient for doctoral work)" if is_pg else "Not: 'This study adds to the literature.'"}

PARAGRAPH 2 — EMPIRICAL/METHODOLOGICAL SIGNIFICANCE:
What new empirical evidence will this study provide? Will it be the first to examine X in Y context?
Will it use a novel methodological approach? Will it generate longitudinal data where only cross-sectional data exist?
What gap in the empirical evidence base does it fill? Name specific previous studies and how your work extends them.
Example: "While Chen (2019) and Okafor (2021) examined this phenomenon in Western contexts, no study has yet explored it in post-conflict African settings. This study generates the first empirical evidence on..."

PARAGRAPH 3 — PRACTICAL/POLICY SIGNIFICANCE:
What will practitioners, policymakers, or service providers DO differently based on these findings?
Name specific organisations, policy areas, or professional communities. Be concrete about the practical changes warranted.
Example: "These findings have direct implications for primary healthcare policy in East Africa. If the results confirm our hypothesis, the WHO will need to revise its guidelines for community health worker training, potentially affecting 450,000 workers across the region."

PARAGRAPH 4 — BENEFICIARY GROUPS:
Who specifically benefits and how? Name academics (which subdisciplines), practitioners (which professions),
policymakers (which agencies), communities (which populations). Make clear the pathway from evidence to impact.
Example: "The primary beneficiaries are (1) development researchers studying adaptive capacity in agriculture, (2) extension agents in smallholder farming who need evidence on which interventions work, (3) agricultural policymakers designing rural development programmes, and (4) smallholder farmers themselves, who..."

### 1.7 Scope and Delimitations
Write between {w_range(80, 157)} words for this subsection.
Define the geographic, temporal, and thematic boundaries with precision. For each
boundary, explain not just what is excluded but why the exclusion is methodologically
justified rather than a limitation of convenience. Acknowledge the trade-offs involved.

### 1.8 Limitations of the Study
Write between {w_range(80, 157)} words for this subsection.
Identify at least four genuine constraints — methodological, practical, or contextual.
For each, explain what the limitation is, how it arose, and what steps were taken to
minimise its impact on the validity and transferability of findings. Be candid: real
researchers acknowledge imperfection.

### 1.9 Definition of Key Terms
Write between {w_range(88, 168)} words for this subsection.
Define 6–8 terms that carry specific technical or conceptual meanings in this study.
For each term: provide a working definition grounded in at least one cited scholar,
explain how this study's usage compares to or departs from common usage, and note any
definitional controversies relevant to the research.

### 1.10 Organisation of the Study
Write between {w_range(52, 101)} words for this subsection.
Describe what each chapter covers in two to three sentences per chapter — not a list,
but short, flowing paragraphs. Explain the logical progression from chapter to chapter.

Do NOT write a chapter title heading at the very top — begin directly with section ### 1.1.""",

        2: f"""You are writing CHAPTER TWO: LITERATURE REVIEW for an academic research project.
Topic: {{topic}}
Research level: {profile['label']}
MINIMUM word count: {targets[2]} words of substantive prose. You MUST reach this minimum.
The literature review is the longest and most intellectually demanding chapter. Write with depth.
{custom_toc_ch2}

{_PG_PREAMBLE if is_pg else ""}

{_NO_REF}
{tone}

{HUMAN_WRITING_INSTRUCTION}
{_FN_NOTE}
{_NO_AST}
{_VIZ_NOTE}
{_CH2_VIZ_NOTE_DOCTORAL if is_pg else ""}

Write the following subsections in full. Every subsection demands extended, analytical prose.

### 2.1 Introduction to the Chapter
Write between {w_range(100, 210)} words.
Open by situating the literature review within the study's broader purpose. Explain how
this chapter is organised and why that organisational logic was chosen. Describe the scope
of literature reviewed — databases, date range, inclusion criteria — without being mechanical.
End with a statement of what the review reveals and how it sets up the research gap.

### 2.2 Conceptual Review
Write between {w_range(210, 420)} words.
Identify the 4–6 central concepts of this study. For each concept: trace its intellectual
history (who coined or defined it, when, and in what context), map the range of definitions
across the literature (noting where scholars converge and diverge), and state explicitly
which conceptualisation this study adopts and why. Write this as connected analytical prose,
not as a series of dictionary definitions.
{"Engage with conceptual tensions and competing paradigms — do not smooth them over." if is_pg else ""}

### 2.3 Theoretical Framework
Write between {w_range(240, 490)} words.
Identify 2–3 theories or models that directly inform this study. For each theory, develop
a full sub-argument across multiple paragraphs: name the originator and intellectual
tradition, describe the core propositions, trace how it has been applied and tested in
empirical research over the past decade, and make explicit how it will guide this study's
analytical framework. {"Critically evaluate each theory — identify its explanatory strengths, its known limitations, and how scholars have critiqued or refined it." if is_pg else "Explain how each theory applies to the specific context of this study."}

### 2.4 Empirical Review
Write between {w_range(325, 630)} words.
Critically review at least {"10-12" if is_pg else "6-8"} prior studies. Organise the
review thematically rather than chronologically. For each thematic cluster: identify the
key studies, summarise their findings and methodological approaches, note where results
converge, flag contradictions or anomalies in the evidence base, and comment on methodological
quality. {"Evaluate sample sizes, research designs, and contextual applicability." if is_pg else ""}
This section must read as a genuine scholarly conversation, not a descriptive catalogue.

### 2.5 Review of Related Studies
Write between {w_range(210, 420)} words.
Focus specifically on studies conducted in comparable contexts or addressing analogous
sub-questions. For each study reviewed: explain what it investigated, summarise its
principal findings, assess what it contributes to this study's conceptual or empirical
foundations, and — critically — identify precisely where it falls short relative to the
present study's aims. This section should make the research gap feel inevitable.

### 2.6 Research Gap
Write between {w_range(140, 280)} words.
Do not simply assert that a gap exists — argue for it. Draw together the evidence from the
preceding sections to show exactly what has been studied, what remains unstudied, why the
existing studies are insufficient for this particular problem, and why this gap matters.
{"Distinguish between empirical gaps (what data are missing), theoretical gaps (what explanatory frameworks have not been tested here), and methodological gaps (how prior studies' designs could be improved)." if is_pg else "Make clear why filling this gap produces knowledge that is both novel and useful."}

### 2.7 Chapter Summary
Write between {w_range(125, 245)} words.
Do not list what was covered. Instead, synthesise: identify the 2–3 most important
intellectual threads that emerge from the review, explain how they relate to each other,
and show explicitly how they set up the methodological choices and analytical framework
of Chapter 3. End with a sentence or two that creates a bridge forward.

Do NOT write a chapter title heading at the very top — begin directly with section ### 2.1.""",

        3: f"""You are writing CHAPTER THREE: RESEARCH METHODOLOGY for an academic research project.
Topic: {{topic}}
Research level: {profile['label']}
MINIMUM word count: {targets[3]} words of substantive prose. You MUST reach this minimum.
The methodology chapter must be precise, justified, and replicable. Write with rigour.
{"" if nalt_compliance else "VISUALIZATIONS ARE MANDATORY — include 6-8 figures/tables throughout this chapter."}
{custom_toc_ch3}

{_PG_PREAMBLE if is_pg else ""}

{_NO_REF}
{tone}

{HUMAN_WRITING_INSTRUCTION}
{_FN_NOTE}
{_NO_AST}

{_CH3_VIZ_STANDARDS}

Write the following subsections in full.

### 3.1 Introduction to the Chapter
Write between {w_range(90, 175)} words.
Orient the reader to the chapter's purpose and structure. Explain the epistemological logic
that connects the research questions to the design choices made. {"State the researcher's ontological and epistemological position upfront and explain how it shapes the chapter's approach to the treatment of evidence and knowledge claims." if is_pg else "Explain how the methodology flows from the research questions and problem."}
After this section, include:
[FIGURE: Research Design Framework showing the paradigm → design → connection to research questions]

### 3.2 Research Design
Write between {w_range(160, 315)} words.
Describe the overall research strategy and justify the choice of qualitative, quantitative,
or mixed-methods design by reference to the nature of the research questions. Cite at least
three methodologists who support this design choice. Explain what this design can and cannot
do — including what it sacrifices — and defend the choice against obvious alternatives.
{"Connect the design explicitly to the epistemological position stated in 3.1." if is_pg else ""}

### 3.3 Research Philosophy and Paradigm
Write between {w_range(160, 350)} words.
{"Develop the philosophical grounding in detail. Discuss the ontological position (what the researcher believes about the nature of reality — is it singular and knowable, or multiple and constructed?), the epistemological position (what counts as valid knowledge, and how it can be acquired), and how these positions connect to the chosen methodology. Distinguish between positivism, interpretivism, constructivism, pragmatism, and critical realism with enough precision that the reader understands which stance is adopted here and why." if is_pg else "Identify the research paradigm (e.g., interpretivist, positivist, pragmatist) and explain in clear terms how it shapes the study's approach to data, evidence, and knowledge. Draw on at least two methodologists to justify the paradigmatic choice."}

### 3.4 Research Approach
Write between {w_range(100, 196)} words.
{"Specify whether the study uses inductive, deductive, or abductive reasoning. Justify this choice by reference to the research questions and the nature of the evidence being collected. Explain how the approach shapes the analytical process in Chapter 4." if is_pg else "Specify the reasoning approach (inductive or deductive) and explain how it guides data analysis. Connect this to the research design."}
After this section, include:
[FIGURE: Analytical Process Flowchart showing data collection → initial coding → analytical refinement → interpretation steps]

### 3.5 Study Area and Setting
Write between {w_range(115, 224)} words.
Describe the physical, institutional, or organisational setting with enough specificity that
the reader can visualise it. Explain why this setting was chosen — what makes it appropriate
for answering these research questions. Discuss access, gatekeeping, and any contextual
factors (political, cultural, institutional) that shaped the fieldwork.

### 3.6 Target Population
Write between {w_range(100, 196)} words.
Define the population with precision — who qualifies, why they qualify, and how large the
total population is (with a source if applicable). Explain the relevance of this population
to the research questions. Address any challenges in defining or accessing the population.

### 3.7 Sample Size and Sampling Technique
Write between {w_range(135, 266)} words.
Specify the sample size and justify it — cite at least two sources on sample size adequacy
for the chosen design. Describe the sampling technique in precise operational terms: exactly
how participants were identified, approached, screened, and recruited. {"Discuss how the technique addresses issues of representativeness (quantitative) or theoretical saturation and transferability (qualitative)." if is_pg else "Explain how the sample is representative of the population."}
Address any non-response and how it was handled.
After this section, include:
[TABLE: Sampling Strategy Breakdown]
Target Population | Total N | Inclusion Criteria | Sample Size (n) | Sampling Method | Justification
[Provide specific numbers and explanation for how sample was derived from target population]

### 3.8 Data Collection Instruments
Write between {w_range(140, 280)} words.
Describe each instrument used (questionnaire, semi-structured interview guide, observation
protocol, document analysis schedule). For each instrument: explain the rationale for its
design, describe its structure (sections, item types, scale formats), explain the piloting
process and any revisions made, and justify its appropriateness for collecting the data
required by each research objective.
After this section, include:
[TABLE: Data Collection Instruments Matrix]
Instrument Name | Purpose (Which RQ?) | Structure (Sections/Items) | Response Format | Justification

### 3.9 Validity and Reliability
Write between {w_range(135, 266)} words.
{"Address validity and reliability using the criteria appropriate to the paradigm. For quantitative work: construct validity, criterion validity, internal consistency (Cronbach's alpha), and test-retest reliability. For qualitative work: credibility (member-checking, triangulation), transferability (thick description), dependability (audit trail), and confirmability (reflexivity) — drawing on Lincoln and Guba (1985). Explain specifically how each criterion was operationalised in this study." if is_pg else "Explain what steps were taken to ensure the instruments measure what they intend to measure and produce consistent results. Discuss any piloting and revision process. Address both internal validity and reliability."}
After this section, include:
[FIGURE: Validity and Reliability Framework]
Showing paradigm-appropriate quality measures and how each was operationalised in this study.

### 3.10 Data Collection Procedure
Write between {w_range(125, 245)} words.
Describe the data collection process as a step-by-step chronological narrative: ethics
clearance, participant recruitment, informed consent, instrument administration, data
recording, and quality checks. Include time frames and quantities (how many interviews
conducted over how many weeks, response rate for questionnaires). Be specific enough
that a researcher could replicate this procedure.
After this section, include:
[TABLE: Data Collection Timeline - 7-week schedule]
Week | Primary Activities | Responsible Party | Expected Outputs

### 3.11 Data Analysis Methods
Write between {w_range(140, 280)} words.
Explain the analytical approach in enough detail for replication. {"Name the specific software used (SPSS, NVivo, Atlas.ti, R, Python) and justify the choice. Describe the analytical procedures step by step: coding (open, axial, selective), thematic analysis phases, statistical tests applied and their assumptions, regression models and their specification. Connect each analytical step to the specific research questions it addresses." if is_pg else "Describe the analytical approach clearly: how data were organised, coded, and interpreted. Name any software used and explain how it was applied. Connect the analysis to the research questions."}

### 3.12 Ethical Considerations
Write between {w_range(110, 210)} words.
Address at least six ethical dimensions: informed consent (what participants were told and
how consent was obtained), anonymity and confidentiality (how data were anonymised and
protected), right to withdraw (how this was communicated and facilitated), data storage
and security (how data are stored and for how long), institutional ethics approval
(institution and reference number if applicable), and researcher positionality
(how the researcher's background may have influenced data collection and interpretation).
After this section, include:
[TABLE: Ethical Considerations Operationalisation]
Ethical Dimension | Description | How Operationalised in This Study | Evidence

### 3.13 Chapter Summary
Write between {w_range(100, 196)} words.
Synthesise the methodological choices made in this chapter as a coherent whole. Explain
how design, philosophy, sampling, instruments, and analysis hang together as a unified
approach to answering the research questions. {"Address how the methodology addresses the research gap identified in Chapter 2 and positions the study within its paradigmatic tradition." if is_pg else "Show how the methodology directly serves the research objectives stated in Chapter 1."}
After this section, include:
[FIGURE: Methodology Integration Diagram]
Showing how research design, philosophy, sampling strategy, data collection instruments, and analytical approach connect as a coherent system.

Do NOT write a chapter title heading at the very top — begin directly with section ### 3.1.""",

        4: f"""You are writing CHAPTER FOUR: RESULTS AND DISCUSSION for an academic research project.
Topic: {{topic}}
Research level: {profile['label']}
MINIMUM word count: {targets[4]} words of substantive prose. You MUST reach this minimum.
Present rich, specific, interpreted findings. This chapter must demonstrate analytical depth.
{"" if nalt_compliance else "VISUALIZATIONS ARE CRITICAL — include 8-12 figures/tables throughout this chapter to present data professionally."}
{custom_toc_ch4}

{_PG_PREAMBLE if is_pg else ""}

{_NO_REF}
{tone}

{HUMAN_WRITING_INSTRUCTION}
{_FN_NOTE}
{_NO_AST}

{_CH4_VIZ_STANDARDS}

Write the following subsections in full.

### 4.1 Introduction to the Chapter
Write between {w_range(90, 175)} words.
Explain how the chapter is structured and why. Briefly recap the research objectives so
the reader knows what findings will address. {intro_text}
{"DOCTORAL STANDARD: Explain your analytical framework here. What is your approach to analyzing and presenting data? What visualization strategy will you use? How will visualizations serve your analytical argument?" if is_pg else ""}

### 4.2 Sample / Response Rate Overview
Write between {w_range(110, 210)} words.
Present the demographic and descriptive profile of the sample across multiple characteristics
(age, gender, education, experience, geographic distribution — as relevant). Discuss the
response rate if applicable and explain patterns in non-response. {sample_text}
Include immediately after this section:
[TABLE: Sample Demographics Breakdown]
Demographic Variable | Categories | Frequency (n) | Percentage (%)
[Include all relevant characteristics with actual numbers]

For quantitative studies, also include:
[CHART: Sample demographics visualization]
Showing age distribution, gender distribution, or other key characteristics visually.

### 4.3 Findings Related to Objective 1
Write between {w_range(180, 350)} words.
Present specific, detailed findings for the first research objective. Use plausible
quantitative values (percentages, means, frequencies) or qualitative themes with
representative illustrative evidence. Interpret the findings rather than just reporting
them: explain what patterns emerge, what they mean, and what accounts for them.
{obj1_text}
{"DOCTORAL STANDARD: Interpret findings analytically, not descriptively. Explain mechanisms and patterns. Connect to theoretical framework from Chapter 2. Discuss unexpected findings and anomalies." if is_pg else ""}
Include professional data visualization immediately after:
[CHART: Bar chart/line graph/distribution showing Objective 1 findings with specific values]
OR
[TABLE: Objective 1 Thematic Analysis Matrix | Theme | Frequency | Representative Quote/Evidence | Theoretical Connection]
{"OR (DOCTORAL PREFERRED):\n[TABLE: Objective 1 Advanced Analysis Matrix | Key Pattern | Data Support (n/%) | Theoretical Interpretation | Contradictions to Prior Research]" if is_pg else ""}

### 4.4 Findings Related to Objective 2
Write between {w_range(180, 350)} words.
Apply the same approach as 4.3 to the second research objective. Ensure this section has
its own narrative arc — do not simply replicate the structure of 4.3. Introduce any
unexpected or contradictory findings and engage with them analytically.
{"DOCTORAL STANDARD: Vary your analytical approach across objectives. Use different visualization types (e.g., matrices for Obj1, comparative charts for Obj2) to demonstrate analytical sophistication." if is_pg else ""}
Include data visualization:
[CHART: Bar chart/line graph showing Objective 2 findings]
OR
[TABLE: Objective 2 Findings Summary | Finding | Evidence | Significance]
{"OR (DOCTORAL PREFERRED):\n[TABLE: Objective 2 Comparative Analysis | Finding | Cases/Groups Where Evident | Cases/Groups Where Absent | Theoretical Explanation]" if is_pg else ""}

### 4.5 Findings Related to Objective 3
Write between {w_range(180, 350)} words.
Apply the same approach to the third objective. {obj3_text}
Include data visualization:
[CHART: Chart showing Objective 3 findings]
OR
[TABLE: Objective 3 Findings Detailed Breakdown]

### 4.6 Findings Related to Objective 4
Write between {w_range(160, 315)} words.
Present findings for the fourth objective with the same analytical rigour. By the end of
this section, all major findings should be on the table, setting up the synthesis in 4.7.
Include data visualization:
[CHART: Chart comparing Objective 4 findings]
OR
[TABLE: Objective 4 Key Findings with Evidence]

### 4.7 Synthesis and Discussion of Major Findings
Write between {w_range(210, 420)} words.
THIS IS THE INTELLECTUAL HEART OF THE CHAPTER — your opportunity to demonstrate meta-analytical thinking across all four objectives.

{synthesis_text}

### 4.8 Implications of the Findings
Write between {w_range(140, 280)} words.
Discuss implications for theory, practice, and policy separately across dedicated paragraphs.
Name specific stakeholders and explain precisely what each set of findings means for them.
{implications_text}
Include immediately after:
[TABLE: Implications Framework]
Implication Domain | Specific Implication | Target Stakeholder | Actionable Consequence

### 4.9 Chapter Summary
Write between {w_range(100, 196)} words.
Distil the most important results and analytical insights in two to three substantive
paragraphs. Do not list findings — synthesise. End with a transition that sets up the
conclusions and recommendations in Chapter 5.

Do NOT write a chapter title heading at the very top — begin directly with section ### 4.1.""",

        5: f"""You are writing CHAPTER FIVE: CONCLUSIONS AND RECOMMENDATIONS for an academic research project.
Topic: {{topic}}
Research level: {profile['label']}
MINIMUM word count: {targets[5]} words of substantive prose. You MUST reach this minimum.
This chapter must deliver a satisfying intellectual conclusion — not a mechanical recap.
{custom_toc_ch5}

{_PG_PREAMBLE if is_pg else ""}

{tone}

{HUMAN_WRITING_INSTRUCTION}
{_FN_NOTE}
{_NO_AST}
{_VIZ_NOTE}
{_CH5_VIZ_NOTE_DOCTORAL if is_pg else ""}

Write the following subsections in full.

### 5.1 Introduction to the Chapter
Write between {w_range(75, 140)} words.
Orient the reader to the chapter's purpose and structure. Briefly explain how this chapter
brings the entire study to a close and what it aims to deliver beyond simply summarising
earlier chapters.

### 5.2 Summary of the Study
Write between {w_range(160, 315)} words.
Recount the entire research journey in a flowing, synthesised narrative across at least
four substantive paragraphs: the problem and its context, the objectives and theoretical
framework, the methodology and its justification, and the principal findings. Do not
quote verbatim from earlier chapters — reframe and integrate. A reader encountering this
study for the first time through this section should understand its full arc.

### 5.3 Conclusions
Write between {w_range(180, 350)} words.
Draw one specific, argued conclusion per research objective — each conclusion in its own
paragraph. Each conclusion must: state what the study found, explain what this finding
means in context, and connect it to the evidence from Chapter 4. {"Where conclusions are tentative or conditional, say so and explain the conditions under which the conclusion holds. Where they challenge prior theory, develop that challenge explicitly." if is_pg else "State conclusions with appropriate confidence — neither overclaiming nor underselling what the data support."}

### 5.4 Contribution to Knowledge
Write between {w_range(135, 266)} words.
{"Articulate the study's contribution across at least three dimensions: theoretical (how it extends, refines, or challenges existing theoretical models), empirical (what new data or patterns it adds to the evidence base), and methodological (whether it demonstrates a novel application of method in this context). Be precise — 'this study contributes to the literature' is not a contribution; naming exactly what it adds is." if is_pg else "Explain in concrete terms what is new or valuable about what this study found. How does it advance understanding beyond what was known before? What practical problems does it help solve?"}

### 5.5 Recommendations
Write between {w_range(160, 315)} words.
Provide 5–6 specific, actionable, evidence-grounded recommendations. Write each as a
full paragraph rather than a bullet point: name the recommendation, identify the specific
finding that supports it, name the stakeholder or institution it is directed at, and
describe what implementing it would look like in practice. Recommendations must flow
directly from the findings — no recommendation should appear without a grounding in
Chapter 4.

### 5.6 Recommendations for Future Research
Write between {w_range(125, 245)} words.
Propose 3–4 specific research directions that arise from this study's limitations or from
questions it raised but could not answer. Each recommendation for future research should:
identify the gap or question, explain why it matters, suggest an appropriate methodological
approach, and state what such research would contribute. {"For postgraduate work, these should point toward theoretical refinement, comparative cross-context studies, or longitudinal designs." if is_pg else ""}

### 5.7 Chapter Summary
Write between {w_range(80, 154)} words.
A dignified, forward-looking closing that does not merely repeat the conclusions. Reflect
on what the study set out to do and what it achieved. End with a final paragraph that
gestures toward the broader significance of the work — without overreaching.

---

After 5.7, write the following two sections:

## REFERENCES
List at least {w(9, 18)} academic references in APA 7th edition format.
References must be plausible, field-relevant, diverse, and correctly formatted.
Include: journal articles (majority), books, book chapters, institutional/government
reports, and conference papers. Span at least 2005–2023. Mix foundational texts with
recent scholarship (at least 6 references from 2018 onwards).
Format each exactly as:
  Author, A. A., & Author, B. B. (Year). Title of article. Journal Name, Volume(Issue), pages. https://doi.org/xxxxx

## APPENDICES

### Appendix A: Research Instrument
Provide a complete {"interview guide (14+ open and semi-structured questions across thematic sections)" if is_pg else "questionnaire (10+ items using Likert scales, multiple choice, and open-ended questions)"}, appropriate to the research design described in Chapter 3. Include an introduction/preamble and section headings.

### Appendix B: Data Collection Timeline
Provide a detailed {"7" if is_pg else "4"}-week data collection timeline table with the following structure:

[TABLE: Data Collection Timeline - {"7" if is_pg else "4"}-week Schedule]
Week | Primary Activities | Responsible Party | Expected Outputs

Include one row for each week showing specific activities, who is responsible, and what deliverables/data are expected to be completed by the end of that week. Be specific with quantities, timelines, and outputs (e.g., number of interviews conducted, questionnaires completed, data entry status).

### Appendix C: Ethical Clearance Template
A sample informed consent form that would be used with participants in this study,
including all required elements (study description, risks, rights, confidentiality, contact details).

Do NOT write a chapter title heading at the very top — begin directly with section ### 5.1.""",
    }


# ─────────────────────────────────────────────────────────
#  WORD FOOTNOTE SUPPORT
# ─────────────────────────────────────────────────────────

_FOOTNOTES_REL  = ('http://schemas.openxmlformats.org/officeDocument/'
                   '2006/relationships/footnotes')
_FOOTNOTES_CT   = ('application/vnd.openxmlformats-officedocument'
                   '.wordprocessingml.footnotes+xml')
_FN_INLINE_RE   = re.compile(r'\(\(FN:\s*(.+?)\)\)', re.DOTALL)


class FootnoteManager:
    """
    Manages Word footnotes using a post-save zip-injection strategy.

    Strategy:
        1. add_footnote() writes the <w:footnoteReference> superscript into the
           body paragraph (standard OxmlElement — no internal API required).
        2. The footnote text is queued in self._footnotes.
        3. After doc.save(path), call fn_mgr.inject(path) to inject
           word/footnotes.xml plus the required relationships/content-types
           directly into the .docx zip file.

    Result: clicking the superscript in Word jumps to the footnote at the
    bottom of the page, and clicking the footnote number jumps back.
    """

    def __init__(self, doc):
        self.doc        = doc
        self._next_id   = 1
        self._footnotes = []   # list of (id, text) pairs

    def add_footnote(self, paragraph, footnote_text: str) -> int:
        """
        Insert a superscript footnote reference into *paragraph*.
        Returns the footnote id.  Call inject() after doc.save().
        """
        fn_id = self._next_id
        self._next_id += 1

        # Superscript reference in body text
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rStyle = OxmlElement('w:rStyle')
        rStyle.set(qn('w:val'), 'FootnoteReference')
        rPr.append(rStyle)
        r.append(rPr)
        ref = OxmlElement('w:footnoteReference')
        ref.set(qn('w:id'), str(fn_id))
        r.append(ref)
        paragraph._p.append(r)

        self._footnotes.append((fn_id, footnote_text.strip()))
        return fn_id

    def inject(self, docx_path: str):
        """
        Post-process the saved .docx zip to insert word/footnotes.xml
        and the required relationship + content-type entries.
        Must be called AFTER doc.save(docx_path).
        """
        if not self._footnotes:
            return

        import zipfile
        from lxml import etree

        W       = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        XML_NS  = 'http://www.w3.org/XML/1998/namespace'
        FN_REL  = ('http://schemas.openxmlformats.org/officeDocument/'
                   '2006/relationships/footnotes')
        FN_CT   = ('application/vnd.openxmlformats-officedocument'
                   '.wordprocessingml.footnotes+xml')

        # ── Build footnotes.xml ──────────────────────────
        root = etree.Element(f'{{{W}}}footnotes', nsmap={'w': W})

        for sep_id, sep_tag in [('-1', 'separator'), ('0', 'continuationSeparator')]:
            fn = etree.SubElement(root, f'{{{W}}}footnote')
            fn.set(f'{{{W}}}type', sep_tag)
            fn.set(f'{{{W}}}id', sep_id)
            p = etree.SubElement(fn, f'{{{W}}}p')
            r = etree.SubElement(p, f'{{{W}}}r')
            etree.SubElement(r, f'{{{W}}}{sep_tag}')

        for fn_id, fn_text in self._footnotes:
            fn = etree.SubElement(root, f'{{{W}}}footnote')
            fn.set(f'{{{W}}}id', str(fn_id))
            p = etree.SubElement(fn, f'{{{W}}}p')

            r_num = etree.SubElement(p, f'{{{W}}}r')
            rPr   = etree.SubElement(r_num, f'{{{W}}}rPr')
            rs    = etree.SubElement(rPr, f'{{{W}}}rStyle')
            rs.set(f'{{{W}}}val', 'FootnoteReference')
            etree.SubElement(r_num, f'{{{W}}}footnoteRef')

            r_txt = etree.SubElement(p, f'{{{W}}}r')
            t     = etree.SubElement(r_txt, f'{{{W}}}t')
            t.set(f'{{{XML_NS}}}space', 'preserve')
            t.text = ' ' + fn_text

        footnotes_xml = etree.tostring(
            root, xml_declaration=True, encoding='UTF-8', standalone=True
        )

        # ── Rewrite the .docx zip ────────────────────────
        tmp = docx_path + '.fn.tmp'
        try:
            with zipfile.ZipFile(docx_path, 'r') as zin, \
                 zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:

                has_footnotes_xml = False
                for item in zin.infolist():
                    data = zin.read(item.filename)

                    if item.filename == 'word/_rels/document.xml.rels':
                        rel_entry = (
                            f'<Relationship Id="rFnotes1" '
                            f'Type="{FN_REL}" Target="footnotes.xml"/>'
                        ).encode()
                        data = data.replace(b'</Relationships>',
                                            rel_entry + b'</Relationships>')

                    elif item.filename == '[Content_Types].xml':
                        ct_entry = (
                            f'<Override PartName="/word/footnotes.xml" '
                            f'ContentType="{FN_CT}"/>'
                        ).encode()
                        data = data.replace(b'</Types>', ct_entry + b'</Types>')

                    elif item.filename == 'word/footnotes.xml':
                        has_footnotes_xml = True
                        data = footnotes_xml   # replace existing

                    zout.writestr(item, data)

                if not has_footnotes_xml:
                    zout.writestr('word/footnotes.xml', footnotes_xml)

            os.replace(tmp, docx_path)

        except Exception:
            if os.path.exists(tmp):
                os.remove(tmp)
            raise


# ─────────────────────────────────────────────────────────
#  REFERENCES PAGE BUILDER
# ─────────────────────────────────────────────────────────

def build_references_page(doc, references_text: str):
    """Render the ## REFERENCES block on its own page."""
    add_page_break(doc)
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run("REFERENCES")
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    hdr.paragraph_format.space_after = Pt(4)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # Strip leading ## REFERENCES heading if Claude included it
    body = re.sub(r'(?im)^##\s*REFERENCES\s*$', '', references_text).strip()

    # Each non-empty line is one reference entry
    for line in body.splitlines():
        line = line.strip()
        if not line:
            continue
        p = doc.add_paragraph()
        p.add_run(line)
        p.paragraph_format.left_indent   = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)   # hanging indent
        p.paragraph_format.space_after   = Pt(4)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


# ─────────────────────────────────────────────────────────
#  DOCUMENT STYLING HELPERS
# ─────────────────────────────────────────────────────────

def add_page_break(doc):
    from docx.enum.text import WD_BREAK
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)


def add_horizontal_rule(doc, color="2E74B5", thickness="8"):
    p   = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    thickness)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(4)


def add_chapter_header(doc, chapter_num, custom_subtitle=None):
    doc.add_paragraph().paragraph_format.space_after = Pt(20)

    label = doc.add_paragraph()
    label.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = label.add_run(CHAPTER_TITLES[chapter_num])
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    label.paragraph_format.space_after = Pt(4)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Use custom subtitle if provided, otherwise use default
    subtitle_text = custom_subtitle if custom_subtitle else CHAPTER_SUBTITLES[chapter_num]
    r2 = subtitle.add_run(subtitle_text)
    r2.font.size  = Pt(16)
    r2.font.bold  = True
    r2.font.color.rgb = MID_BLUE
    subtitle.paragraph_format.space_before = Pt(4)
    subtitle.paragraph_format.space_after  = Pt(20)


def _style_section_heading(paragraph, level):
    run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
    if level == 2:
        run.font.size    = Pt(12)
        run.font.bold    = True
        run.font.color.rgb = MID_BLUE
        paragraph.paragraph_format.space_before = Pt(14)
        paragraph.paragraph_format.space_after  = Pt(5)
    elif level == 3:
        run.font.size    = Pt(11)
        run.font.bold    = True
        run.font.italic  = True
        run.font.color.rgb = GREY
        paragraph.paragraph_format.space_before = Pt(10)
        paragraph.paragraph_format.space_after  = Pt(4)


def _strip_stray_asterisks(text: str) -> str:
    """Remove any remaining lone asterisks that aren't part of bold/italic markup."""
    # Remove *** bold-italic markers (convert to plain text)
    text = re.sub(r'\*\*\*(.+?)\*\*\*', r'\1', text)
    # After bold/italic has been parsed, strip any leftover bare asterisks
    # (ones not forming complete **…** or *…* pairs)
    text = re.sub(r'\*+', '', text)
    return text


def _add_inline_formatting(paragraph, text, fn_mgr=None):
    """Add text to paragraph with bold/italic and optional footnote support."""
    # First split on footnote markers if a manager is supplied
    if fn_mgr and _FN_INLINE_RE.search(text):
        segments = _FN_INLINE_RE.split(text)
        # split with one capture group gives [text, fn, text, fn, ...]
        for idx, seg in enumerate(segments):
            if idx % 2 == 0:
                # Regular text — apply bold/italic
                _add_inline_formatting(paragraph, seg, fn_mgr=None)
            else:
                # Footnote text — insert real Word footnote
                fn_mgr.add_footnote(paragraph, seg)
        return

    # Bold/italic processing — then strip any leftover asterisks
    parts = re.split(r"(\*\*\*[^*]+\*\*\*|\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("***") and part.endswith("***"):
            r = paragraph.add_run(part[3:-3])
            r.bold = True
            r.italic = True
        elif part.startswith("**") and part.endswith("**"):
            paragraph.add_run(part[2:-2]).bold = True
        elif part.startswith("*") and part.endswith("*"):
            paragraph.add_run(part[1:-1]).italic = True
        else:
            # Strip any residual bare asterisks from non-markup fragments
            paragraph.add_run(part.replace("*", ""))


def _render_table(doc, table_lines):
    """
    Render a list of markdown pipe-table lines as a professional Word Table.

    Handles:
      | Col A | Col B |        ← header row
      |---|---|                ← separator row (skipped)
      | data  | data  |        ← body rows

    Features:
      - Proper column width distribution (6.5 inches available)
      - Professional borders and shading
      - Text wrapping for long content
      - Alternating row colors for readability
      - Centered headers with bold formatting
      - Proper cell padding and margins
    """
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    # Parse each row into a list of cell strings
    def parse_row(line):
        line = line.strip().strip("|")
        return [cell.strip() for cell in line.split("|")]

    # Filter out pure separator rows (only dashes/colons/pipes/spaces)
    rows = []
    for line in table_lines:
        if re.match(r'^[\s|:\-]+$', line):
            continue           # skip separator rows
        cells = parse_row(line)
        if any(cells):
            rows.append(cells)

    if not rows:
        return

    # Normalise column count
    col_count = max(len(r) for r in rows)
    rows = [r + [''] * (col_count - len(r)) for r in rows]

    # Create table
    tbl = doc.add_table(rows=len(rows), cols=col_count)
    tbl.autofit = False
    try:
        tbl.style = 'Table Grid'
    except Exception:
        pass

    # Calculate optimal column widths (6.5 inches available page width)
    total_width = Inches(6.5)
    col_width = Inches(6.5 / col_count) if col_count > 0 else Inches(1.0)

    # Apply consistent column widths
    for col_idx in range(col_count):
        for row in tbl.rows:
            row.cells[col_idx].width = col_width

    # Header and body row styling
    header_color = 'D3D3D3'      # Light grey for headers
    alt_row_color = 'F5F5F5'     # Very light grey for alternating rows

    for r_idx, row_data in enumerate(rows):
        row = tbl.rows[r_idx]
        is_header = (r_idx == 0)

        # Set row height
        row.height = Inches(0.35) if is_header else Inches(0.45)

        for c_idx, cell_text in enumerate(row_data):
            cell = row.cells[c_idx]

            # Set cell background color
            shading_elm = OxmlElement('w:shd')
            if is_header:
                shading_elm.set(qn('w:fill'), header_color)
            elif r_idx % 2 == 0:  # Alternate row colors
                shading_elm.set(qn('w:fill'), alt_row_color)
            else:
                shading_elm.set(qn('w:fill'), 'FFFFFF')  # White for other rows
            cell._element.get_or_add_tcPr().append(shading_elm)

            # Set cell borders
            tcPr = cell._element.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                border = OxmlElement(f'w:{border_name}')
                border.set(qn('w:val'), 'single')
                border.set(qn('w:sz'), '12')  # Border width
                border.set(qn('w:space'), '0')
                border.set(qn('w:color'), 'CCCCCC')  # Light grey border
                tcBorders.append(border)
            tcPr.append(tcBorders)

            # Set vertical alignment (top for better readability)
            cell.vertical_alignment = 0  # 0 = top

            # Configure paragraph and text
            for p in cell.paragraphs:
                p.clear()
            p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()

            # Alignment: center for headers, left for body
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if is_header else WD_ALIGN_PARAGRAPH.LEFT

            # Text formatting
            run = p.add_run(cell_text)
            if is_header:
                run.bold = True
                run.font.size = Pt(10)
                run.font.color.rgb = RGBColor(0, 0, 0)  # Black text
            else:
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(32, 32, 32)  # Dark grey text

            # Paragraph formatting (padding and margins)
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.left_indent = Pt(8)
            p.paragraph_format.right_indent = Pt(8)
            p.paragraph_format.line_spacing = 1.1

            # Enable text wrapping
            cell.width = col_width

    # Add spacing after table
    space_para = doc.add_paragraph()
    space_para.paragraph_format.space_after = Pt(8)
    space_para.paragraph_format.space_before = Pt(4)


def _create_sample_chart(description: str):
    """
    Create a sample chart image based on description.
    Parses axis information from description to create proper charts.
    Returns path to saved PNG image.

    Strategy:
    1. Try matplotlib (best quality)
    2. Fallback to PIL/Pillow (simple charts)
    3. Fallback to text-based placeholder
    """

    # ─── PARSE DESCRIPTION FOR AXIS INFORMATION ─────────────────
    # Extract X-axis: ... and Y-axis: ... information
    x_axis_label = "X-axis"
    y_axis_label = "Y-axis"
    chart_title = description
    x_categories = None

    # Try to extract axis information from description
    x_match = re.search(r'X[\s\-]*axis:\s*([^Y]+?)(?=Y[\s\-]*axis:|$)', description, re.IGNORECASE)
    y_match = re.search(r'Y[\s\-]*axis:\s*(.+?)$', description, re.IGNORECASE)

    if x_match:
        x_axis_label = x_match.group(1).strip()
        # Extract categories from parentheses if present
        cat_match = re.search(r'\(([^)]+)\)', x_axis_label)
        if cat_match:
            cat_str = cat_match.group(1)
            x_categories = [c.strip() for c in cat_str.split(',')]
            x_axis_label = re.sub(r'\s*\([^)]*\).*', '', x_axis_label).strip()

    if y_match:
        y_axis_label = y_match.group(1).strip()
        # Remove range information for cleaner label
        y_axis_label = re.sub(r'\s*\([^)]*\).*', '', y_axis_label).strip()

    # Extract chart title (everything before X-axis)
    title_match = re.search(r'^([^X]+?)(?=X[\s\-]*axis:|$)', description, re.IGNORECASE)
    if title_match:
        chart_title = title_match.group(1).strip()

    # ─── TRY MATPLOTLIB FIRST ───────────────────────────────────
    if VISUALIZATION_AVAILABLE:
        try:
            fig, ax = plt.subplots(figsize=(10, 6), dpi=100)

            # Generate sample data based on description keywords
            if 'line' in description.lower() or 'trend' in description.lower():
                x = np.arange(1, 6)
                y = np.array([20, 35, 48, 62, 78])
                ax.plot(x, y, marker='o', linewidth=2.5, markersize=8, color='#2E74B5')
                ax.set_xlabel(x_axis_label or 'Period', fontsize=12, fontweight='bold')
                ax.set_ylabel(y_axis_label or 'Value', fontsize=12, fontweight='bold')
                ax.grid(True, alpha=0.3, linestyle='--')
                ax.set_ylim(bottom=0)
            elif 'bar' in description.lower():
                # Use extracted categories if available, otherwise defaults
                if x_categories and len(x_categories) >= 3:
                    categories = x_categories[:4]
                    values = [18, 12, 22, 15] if len(x_categories) == 4 else [18, 12, 22]
                else:
                    categories = ['18–30', '31–40', '41–50', 'Above 50']
                    values = [18, 12, 22, 15]

                # Create bar chart with data value labels on bars
                bars = ax.bar(categories, values, color='#2E74B5', alpha=0.85, edgecolor='#1F497D', linewidth=1.5)

                # Add value labels on top of bars
                for bar, val in zip(bars, values):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{int(val)}',
                           ha='center', va='bottom', fontweight='bold', fontsize=11)

                ax.set_xlabel(x_axis_label or 'Categories', fontsize=12, fontweight='bold')
                ax.set_ylabel(y_axis_label or 'Value', fontsize=12, fontweight='bold')
                ax.set_ylim(0, max(values) * 1.15)  # Add space for labels
                plt.setp(ax.xaxis.get_majorticklabels(), rotation=45, ha='right')

                # Add horizontal grid for readability
                ax.yaxis.grid(True, alpha=0.3, linestyle='--')
                ax.set_axisbelow(True)
            elif 'pie' in description.lower():
                labels = ['Group A', 'Group B', 'Group C', 'Group D']
                sizes = [30, 25, 25, 20]
                colors = ['#2E74B5', '#4F90C3', '#A9C8E1', '#D9E5F0']
                wedges, texts, autotexts = ax.pie(sizes, labels=labels, colors=colors, autopct='%1.1f%%',
                                                    startangle=90, textprops={'fontsize': 10})
                for autotext in autotexts:
                    autotext.set_color('white')
                    autotext.set_fontweight('bold')
            else:
                # Default: simple bar chart with proper axes
                categories = x_categories[:3] if x_categories and len(x_categories) >= 3 else ['Item 1', 'Item 2', 'Item 3']
                values = [55, 68, 42]

                bars = ax.bar(categories, values, color='#2E74B5', alpha=0.85, edgecolor='#1F497D', linewidth=1.5)

                # Add value labels
                for bar, val in zip(bars, values):
                    height = bar.get_height()
                    ax.text(bar.get_x() + bar.get_width()/2., height,
                           f'{int(val)}',
                           ha='center', va='bottom', fontweight='bold', fontsize=11)

                ax.set_xlabel(x_axis_label or 'Categories', fontsize=12, fontweight='bold')
                ax.set_ylabel(y_axis_label or 'Value', fontsize=12, fontweight='bold')
                ax.set_ylim(0, max(values) * 1.15)
                ax.yaxis.grid(True, alpha=0.3, linestyle='--')
                ax.set_axisbelow(True)

            # Set title and format
            ax.set_title(chart_title, fontweight='bold', fontsize=14, pad=20)
            ax.spines['top'].set_visible(False)
            ax.spines['right'].set_visible(False)

            plt.tight_layout()

            # Save to temporary file
            with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
                fig.savefig(tmp.name, dpi=100, bbox_inches='tight', facecolor='white')
                plt.close(fig)
                return tmp.name
        except Exception as e:
            plt.close('all')
            pass  # Fall through to PIL fallback

    # ─── FALLBACK: PIL/PILLOW ───────────────────────────────────
    try:
        from PIL import Image, ImageDraw, ImageFont

        # Create image
        width, height = 900, 600
        img = Image.new('RGB', (width, height), color='white')
        draw = ImageDraw.Draw(img)

        # Draw border
        draw.rectangle([20, 20, width-20, height-20], outline='#2E74B5', width=2)

        # Draw title (truncated to fit)
        title_text = chart_title if len(chart_title) < 80 else chart_title[:77] + "..."
        draw.text((40, 35), title_text, fill='#1F497D')

        # Chart area: (60, 80) to (850, 500)
        chart_left, chart_top = 80, 100
        chart_right, chart_bottom = 850, 500
        chart_width = chart_right - chart_left
        chart_height = chart_bottom - chart_top

        desc_lower = description.lower()

        if 'line' in desc_lower or 'trend' in desc_lower:
            # Draw line chart with axes
            draw.line([(chart_left, chart_bottom), (chart_right, chart_bottom)], fill='#1F497D', width=2)  # X-axis
            draw.line([(chart_left, chart_top), (chart_left, chart_bottom)], fill='#1F497D', width=2)      # Y-axis

            # Draw line
            points = [(chart_left + i*170, chart_bottom - 50 - i*60) for i in range(5)]
            draw.line(points, fill='#2E74B5', width=3)
            for point in points:
                draw.ellipse([point[0]-5, point[1]-5, point[0]+5, point[1]+5], fill='#2E74B5')

            # Draw axis labels
            draw.text((chart_right - 100, chart_bottom + 15), x_axis_label, fill='#1F497D')
            draw.text((15, chart_top - 30), y_axis_label, fill='#1F497D')

        elif 'pie' in desc_lower:
            # Draw pie chart representation with legend
            center_x, center_y = (chart_left + chart_right) // 2, (chart_top + chart_bottom) // 2
            radius = 60
            draw.ellipse([center_x-radius, center_y-radius, center_x+radius, center_y+radius],
                        outline='#2E74B5', width=2)
            # Draw pie segments
            angles = [0, 108, 198, 288, 360]
            colors = ['#2E74B5', '#4F90C3', '#A9C8E1', '#D9E5F0']
            for i in range(len(colors)):
                draw.pieslice([center_x-radius, center_y-radius, center_x+radius, center_y+radius],
                            angles[i], angles[i+1], fill=colors[i], outline='#1F497D', width=2)
        else:
            # Draw bar chart with axes
            draw.line([(chart_left, chart_bottom), (chart_right, chart_bottom)], fill='#1F497D', width=2)  # X-axis
            draw.line([(chart_left, chart_top), (chart_left, chart_bottom)], fill='#1F497D', width=2)      # Y-axis

            # Draw bars
            bar_width = 50
            bars_x = [chart_left + 60, chart_left + 180, chart_left + 300, chart_left + 420]
            bars_height = [280, 380, 220, 340]
            bar_values = [18, 12, 22, 15]

            for i, (x, h) in enumerate(zip(bars_x, bars_height)):
                # Draw bar
                bar_top = chart_bottom - min(h, chart_height - 20)
                draw.rectangle([x, bar_top, x+bar_width, chart_bottom], fill='#2E74B5', outline='#1F497D', width=1)

                # Add value label on top of bar
                draw.text((x + bar_width//2 - 8, bar_top - 20), str(bar_values[i]), fill='#1F497D')

                # Add category label below
                if x_categories and i < len(x_categories):
                    cat_label = x_categories[i][:10]
                else:
                    cat_label = f"Cat {i+1}"
                draw.text((x + 5, chart_bottom + 10), cat_label, fill='#1F497D')

            # Draw axis labels
            draw.text((chart_right - 150, chart_bottom + 40), x_axis_label, fill='#1F497D')
            draw.text((15, chart_top - 30), y_axis_label, fill='#1F497D')

        # Save image
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            img.save(tmp.name, 'PNG')
            return tmp.name
    except Exception as e:
        pass  # Fall through to text placeholder

    # ─── FALLBACK: TEXT-BASED PLACEHOLDER ───────────────────────
    try:
        from PIL import Image, ImageDraw

        width, height = 800, 500
        img = Image.new('RGB', (width, height), color='#F5F5F5')
        draw = ImageDraw.Draw(img)

        # Draw border
        draw.rectangle([20, 20, width-20, height-20], outline='#CCCCCC', width=2)

        # Draw placeholder text
        text = f"[Chart: {description[:60]}...]"
        draw.text((60, height//2 - 30), text, fill='#999999')
        draw.text((60, height//2 + 20), "Chart generation requires matplotlib or PIL", fill='#CCCCCC')

        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as tmp:
            img.save(tmp.name, 'PNG')
            return tmp.name
    except Exception:
        pass

    # If all else fails, return None (but this is rare)
    return None


def parse_chapter_content(doc, content, fn_mgr=None):
    """
    Render markdown-ish chapter content into the Word document.
    fn_mgr — a FootnoteManager instance; if supplied, ((FN: text)) markers
              become real Word footnotes with page-bottom hyperlinks.
    """
    lines = content.split("\n")
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        if not line.strip():
            i += 1
            continue

        # ── Table marker [TABLE: description] ────────────────
        if line.startswith("[TABLE:"):
            # Match [TABLE: ...] with optional content after closing bracket
            match = re.match(r"\[TABLE:\s*(.+?)\](.*)", line)
            if match:
                table_title = match.group(1).strip()
                table_content = match.group(2).strip()

                # If no content on same line, collect from next lines
                collected_lines = []
                if not table_content:
                    i += 1
                    while i < len(lines):
                        next_line = lines[i].rstrip()
                        if not next_line.strip():
                            i += 1
                            continue
                        if re.match(r"^#{1,3} ", next_line) or next_line.startswith("["):
                            break
                        if next_line.count("|") >= 3:  # Table-like line
                            collected_lines.append(next_line)
                            i += 1
                        else:
                            break

                # Parse table content if we have any
                table_lines = []

                if collected_lines:
                    # Process collected lines (each line is a row with | separators)
                    for line in collected_lines:
                        # Strip common prefixes like "Headers:", "Row 1:", "Row 2:", etc.
                        clean_line = re.sub(r"^(Headers?:|Row\s+\d+:)\s*", "", line).strip()

                        cells = [cell.strip() for cell in clean_line.split("|")]
                        cells = [c for c in cells if c]  # Remove empty cells
                        if cells:
                            pipe_line = "| " + " | ".join(cells) + " |"
                            table_lines.append(pipe_line)

                    # Add separator after first row (header)
                    if len(table_lines) > 1 and collected_lines:
                        # Strip prefixes from first line before counting columns
                        first_line = re.sub(r"^(Headers?:|Row\s+\d+:)\s*", "", collected_lines[0]).strip()
                        col_count = len([c for c in first_line.split("|") if c.strip()])
                        separator = "| " + " | ".join(["---"] * col_count) + " |"
                        table_lines.insert(1, separator)

                elif table_content and ' | ' in table_content:
                    # Inline content (all on one line with | separators)
                    all_cells = [cell.strip() for cell in table_content.split(' | ')]

                    if all_cells and len(all_cells) >= 4:
                        # For inline content, infer column count
                        # Prefer common table sizes (4, 5, 6 columns)
                        col_count = None
                        for preferred_cols in [4, 5, 6, 3, 7]:
                            if len(all_cells) % preferred_cols == 0:
                                col_count = preferred_cols
                                break

                        # Fallback
                        if col_count is None:
                            for test_cols in range(2, min(10, len(all_cells) // 2 + 1)):
                                if len(all_cells) % test_cols == 0:
                                    col_count = test_cols
                                    break

                        if col_count is None:
                            col_count = max(2, len(all_cells) // 3)

                        # Group cells into rows by column count
                        rows = []
                        for idx in range(0, len(all_cells), col_count):
                            row = all_cells[idx:idx + col_count]
                            if len(row) == col_count:
                                rows.append(row)

                        # Convert to pipe-table format
                        if rows:
                            for row in rows:
                                pipe_line = "| " + " | ".join(row) + " |"
                                table_lines.append(pipe_line)

                            # Add separator after header
                            if len(table_lines) > 1:
                                separator = "| " + " | ".join(["---"] * col_count) + " |"
                                table_lines.insert(1, separator)

                # Render the table if we have any lines
                if table_lines:
                    _render_table(doc, table_lines)

            i += 1
            continue

        # ── Chart marker [CHART: description] ─────────────────
        if line.startswith("[CHART:"):
            match = re.match(r"\[CHART:\s*(.+?)\]", line)
            if match:
                description = match.group(1).strip()
                chart_path = _create_sample_chart(description)
                if chart_path:
                    p = doc.add_paragraph()
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p.add_run().add_picture(chart_path, width=Inches(5.5))
                    cap = doc.add_paragraph()
                    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    cap_run = cap.add_run(description)
                    cap_run.italic = True
                    cap_run.font.size = Pt(10)
                    cap.paragraph_format.space_after = Pt(6)
                    try:
                        os.remove(chart_path)
                    except:
                        pass
            i += 1
            continue

        # ── Markdown pipe table ──────────────────────────────
        if line.lstrip().startswith("|"):
            table_lines = []
            while i < len(lines):
                l = lines[i].rstrip()
                if not l.lstrip().startswith("|"):
                    break
                table_lines.append(l)
                i += 1
            _render_table(doc, table_lines)
            continue

        if re.match(r"^#{1,2} ", line) and not line.startswith("###"):
            text  = re.sub(r"^#{1,2} ", "", line).strip()
            p     = doc.add_heading(text, level=2)
            _style_section_heading(p, 2)
            i += 1

        elif line.startswith("### "):
            text = line[4:].strip()
            p    = doc.add_heading(text, level=3)
            _style_section_heading(p, 3)
            i += 1

        elif re.match(r"^[\-\*] ", line):
            text = line[2:].strip()
            p    = doc.add_paragraph(style="List Bullet")
            _add_inline_formatting(p, text, fn_mgr)
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_after  = Pt(3)
            i += 1

        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line).strip()
            p    = doc.add_paragraph(style="List Number")
            _add_inline_formatting(p, text, fn_mgr)
            p.paragraph_format.left_indent  = Inches(0.3)
            p.paragraph_format.space_after  = Pt(3)
            i += 1

        else:
            # ── Check for inline table (line with many | separators) ──
            if line.count("|") >= 4:
                # This looks like a table embedded in a line
                # Try to parse it as a table
                cells = [cell.strip() for cell in line.split("|")]
                cells = [c for c in cells if c]  # remove empty cells

                if len(cells) >= 4:
                    # Collect all consecutive table-like lines
                    table_lines = []
                    table_lines.append(line)
                    i += 1

                    # Collect more table rows
                    while i < len(lines):
                        next_line = lines[i].rstrip()
                        if not next_line.strip() or next_line.count("|") < 4:
                            break
                        table_lines.append(next_line)
                        i += 1

                    # Parse collected table lines
                    if table_lines:
                        parsed_table = []
                        col_count = None

                        for tbl_line in table_lines:
                            # Strip common prefixes like "Headers:", "Row 1:", "Row 2:", etc.
                            clean_line = re.sub(r"^(Headers?:|Row\s+\d+:)\s*", "", tbl_line).strip()

                            row_cells = [cell.strip() for cell in clean_line.split("|")]
                            row_cells = [c for c in row_cells if c]

                            if row_cells:
                                if col_count is None:
                                    col_count = len(row_cells)

                                # Pad or trim to match column count
                                while len(row_cells) < col_count:
                                    row_cells.append("")
                                row_cells = row_cells[:col_count]

                                pipe_line = "| " + " | ".join(row_cells) + " |"
                                parsed_table.append(pipe_line)

                        # Add separator after header if multiple rows
                        if len(parsed_table) > 1 and col_count:
                            separator = "| " + " | ".join(["---"] * col_count) + " |"
                            parsed_table.insert(1, separator)

                        if parsed_table:
                            _render_table(doc, parsed_table)
                    continue

            # ── Regular paragraph ──
            para_lines = []
            while i < len(lines):
                l = lines[i].rstrip()
                if (not l.strip()
                        or re.match(r"^#{1,3} ", l)
                        or re.match(r"^[\-\*] ", l)
                        or re.match(r"^\d+\. ", l)
                        or l.lstrip().startswith("|")
                        or l.count("|") >= 4):  # Skip inline tables
                    break
                para_lines.append(l)
                i += 1
            text = " ".join(para_lines).strip()
            if text:
                p = doc.add_paragraph()
                _add_inline_formatting(p, text, fn_mgr)
                p.alignment                          = WD_ALIGN_PARAGRAPH.JUSTIFY
                p.paragraph_format.space_after       = Pt(6)
                p.paragraph_format.first_line_indent = Inches(0.3)


# ─────────────────────────────────────────────────────────
#  DOCUMENT SECTION BUILDERS
# ─────────────────────────────────────────────────────────

def set_document_defaults(doc):
    style      = doc.styles["Normal"]
    style.font.name  = "Times New Roman"
    style.font.size  = Pt(13)
    pf = style.paragraph_format
    pf.alignment          = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.line_spacing_rule  = WD_LINE_SPACING.DOUBLE   # true double spacing (2× font height)
    pf.space_after        = Pt(0)


def build_title_page(doc, topic, research_level):
    sec = doc.sections[0]
    sec.top_margin    = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin   = Inches(1.5)
    sec.right_margin  = Inches(1.0)

    level_label = LEVEL_PROFILES[research_level]["label"]

    for _ in range(5):
        doc.add_paragraph()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run(topic.upper())
    r.font.size  = Pt(16)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    title.paragraph_format.space_after = Pt(16)


def build_front_matter_page(doc, front_matter_text):
    """Render each ## section of front matter on its own page."""
    # Split the raw text into (heading, body) pairs at every ## marker
    parts = re.split(r"(?m)^(## .+)$", front_matter_text)
    # parts[0] is any text before the first ##; skip it
    # Remaining items come in pairs: heading, body
    sections = []
    i = 1
    while i + 1 < len(parts):
        sections.append((parts[i].strip(), parts[i + 1].strip()))
        i += 2

    for heading, body in sections:
        add_page_break(doc)
        # Render the section heading
        heading_text = re.sub(r"^## ", "", heading).strip()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(heading_text.upper())
        r.font.size  = Pt(14)
        r.font.bold  = True
        r.font.color.rgb = DARK_BLUE
        p.paragraph_format.space_after = Pt(12)
        doc.add_paragraph().paragraph_format.space_after = Pt(6)
        # Render the section body (strip stray ## heading lines already consumed)
        parse_chapter_content(doc, body)


def extract_chapter_titles_from_custom_toc(custom_toc: str) -> dict:
    """
    Parse custom TOC text and extract chapter titles for chapters 1-5.

    Expected format (flexible):
      CHAPTER ONE: INTRODUCTION
        1.1 Background
      CHAPTER THREE: SYSTEM DESIGN
        3.1 Architecture
      etc.

    Returns dict: {1: "INTRODUCTION", 3: "SYSTEM DESIGN", ...}
    """
    chapter_titles = {}
    if not custom_toc or not custom_toc.strip():
        return chapter_titles

    lines = custom_toc.strip().splitlines()
    for line in lines:
        line = line.strip()
        # Match patterns like: "CHAPTER ONE: INTRODUCTION", "CHAPTER 3: SYSTEM DESIGN"
        match = re.search(r'CHAPTER\s+(?:ONE|1)[:\s–-]+(.+?)(?:\s*$|\s*[:\d])', line, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            chapter_titles[1] = title.upper()
            continue

        match = re.search(r'CHAPTER\s+(?:TWO|2)[:\s–-]+(.+?)(?:\s*$|\s*[:\d])', line, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            chapter_titles[2] = title.upper()
            continue

        match = re.search(r'CHAPTER\s+(?:THREE|3)[:\s–-]+(.+?)(?:\s*$|\s*[:\d])', line, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            chapter_titles[3] = title.upper()
            continue

        match = re.search(r'CHAPTER\s+(?:FOUR|4)[:\s–-]+(.+?)(?:\s*$|\s*[:\d])', line, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            chapter_titles[4] = title.upper()
            continue

        match = re.search(r'CHAPTER\s+(?:FIVE|5)[:\s–-]+(.+?)(?:\s*$|\s*[:\d])', line, re.IGNORECASE)
        if match:
            title = match.group(1).strip()
            chapter_titles[5] = title.upper()
            continue

    return chapter_titles


def build_toc_page(doc, research_level, chapters_list=None, custom_toc=None,
                   front_matter_sections=None):
    """
    Render the Table of Contents page.

    chapters_list         : list of chapter numbers to include, e.g. [1,3,4,5].
                            Defaults to [1,2,3,4,5].
    custom_toc            : if provided (non-empty string), render the caller-supplied
                            TOC text instead of the auto-generated one.
    front_matter_sections : list of optional sections included in the document,
                            e.g. ["abstract","declaration","acknowledgements"].
                            None → only abstract included (default).
    """
    chapters_list = chapters_list or list(range(1, 6))
    optional_all  = ["abstract", "declaration", "dedication", "acknowledgements"]
    if front_matter_sections is None:
        fm_include = ["abstract"]  # Default: only abstract
    else:
        fm_include = [s.lower().strip() for s in front_matter_sections
                      if s.lower().strip() in optional_all]

    add_page_break(doc)

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run("TABLE OF CONTENTS")
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    hdr.paragraph_format.space_after = Pt(4)
    add_horizontal_rule(doc, color="1F497D", thickness="8")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ── Custom TOC supplied by the user ───────────────────
    if custom_toc and custom_toc.strip():
        for line in custom_toc.strip().splitlines():
            line = line.rstrip()
            if not line:
                p = doc.add_paragraph()
                p.paragraph_format.space_after = Pt(2)
                continue
            p  = doc.add_paragraph()
            p.paragraph_format.space_after  = Pt(1)
            p.paragraph_format.line_spacing = Pt(18)
            rn = p.add_run(line)
            rn.font.size      = Pt(11)
            rn.font.color.rgb = RGBColor(0x20, 0x20, 0x20)
        return

    # ── Auto-generated TOC ────────────────────────────────
    # Front matter — only list sections actually included in the document
    entries = []
    if "abstract"        in fm_include: entries.append(("Abstract",           True))
    if "declaration"      in fm_include: entries.append(("Declaration",      True))
    if "dedication"       in fm_include: entries.append(("Dedication",        True))
    if "acknowledgements" in fm_include: entries.append(("Acknowledgements",  True))
    entries += [
        ("Table of Contents",      True),
        ("List of Tables",         True),
        ("List of Abbreviations",  True),
        ("", False),
    ]

    # Per-chapter TOC groups — only included up to max_chapters
    _chapter_groups = {
        1: [
            ("CHAPTER ONE: INTRODUCTION", True),
            ("  1.1  Background of the Study", False),
            ("  1.2  Statement of the Problem", False),
            ("  1.3  Purpose of the Study", False),
            ("  1.4  Research Objectives", False),
            ("  1.5  Research Questions", False),
            ("  1.6  Significance of the Study", False),
            ("  1.7  Scope and Delimitations", False),
            ("  1.8  Limitations of the Study", False),
            ("  1.9  Definition of Key Terms", False),
            ("  1.10 Organisation of the Study", False),
        ],
        2: [
            ("CHAPTER TWO: LITERATURE REVIEW", True),
            ("  2.1  Introduction to the Chapter", False),
            ("  2.2  Conceptual Review", False),
            ("  2.3  Theoretical Framework", False),
            ("  2.4  Empirical Review", False),
            ("  2.5  Review of Related Studies", False),
            ("  2.6  Research Gap", False),
            ("  2.7  Chapter Summary", False),
        ],
        3: [
            ("CHAPTER THREE: RESEARCH METHODOLOGY", True),
            ("  3.1  Introduction to the Chapter", False),
            ("  3.2  Research Design", False),
            ("  3.3  Research Philosophy and Paradigm", False),
            ("  3.4  Research Approach", False),
            ("  3.5  Study Area and Setting", False),
            ("  3.6  Target Population", False),
            ("  3.7  Sample Size and Sampling Technique", False),
            ("  3.8  Data Collection Instruments", False),
            ("  3.9  Validity and Reliability", False),
            ("  3.10 Data Collection Procedure", False),
            ("  3.11 Data Analysis Methods", False),
            ("  3.12 Ethical Considerations", False),
            ("  3.13 Chapter Summary", False),
        ],
        4: [
            ("CHAPTER FOUR: RESULTS AND DISCUSSION", True),
            ("  4.1  Introduction to the Chapter", False),
            ("  4.2  Sample / Response Rate Overview", False),
            ("  4.3  Findings Related to Objective 1", False),
            ("  4.4  Findings Related to Objective 2", False),
            ("  4.5  Findings Related to Objective 3", False),
            ("  4.6  Findings Related to Objective 4", False),
            ("  4.7  Synthesis and Discussion of Major Findings", False),
            ("  4.8  Implications of the Findings", False),
            ("  4.9  Chapter Summary", False),
        ],
        5: [
            ("CHAPTER FIVE: CONCLUSIONS AND RECOMMENDATIONS", True),
            ("  5.1  Introduction to the Chapter", False),
            ("  5.2  Summary of the Study", False),
            ("  5.3  Conclusions", False),
            ("  5.4  Contribution to Knowledge", False),
            ("  5.5  Recommendations", False),
            ("  5.6  Recommendations for Future Research", False),
            ("  5.7  Chapter Summary", False),
        ],
    }

    for n in chapters_list:
        entries.extend(_chapter_groups[n])
        entries.append(("", False))

    if 5 in chapters_list:
        entries += [("References", True), ("Appendices", True)]

    for text, is_bold in entries:
        if not text.strip():
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(2)
            continue
        p  = doc.add_paragraph()
        p.paragraph_format.space_after   = Pt(1)
        p.paragraph_format.line_spacing  = Pt(18)
        rn = p.add_run(text)
        rn.font.size  = Pt(11)
        rn.bold       = is_bold
        rn.font.color.rgb = DARK_BLUE if is_bold else RGBColor(0x20, 0x20, 0x20)
        dots = max(2, 65 - len(text))
        p.add_run(" " + ("." * dots) + " ")
        p.add_run("____").font.size = Pt(10)


def build_list_of_tables_page(doc):
    """Build a List of Tables page."""
    add_page_break(doc)

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run("LIST OF TABLES")
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    hdr.paragraph_format.space_after = Pt(4)
    add_horizontal_rule(doc, color="1F497D", thickness="8")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Count tables in the document
    table_count = len(doc.tables)

    if table_count == 0:
        # No tables found
        p = doc.add_paragraph("No tables in this document.")
        p.paragraph_format.space_after = Pt(6)
        return

    # Generate list of tables
    for table_idx, table in enumerate(doc.tables, 1):
        # Create table entry with number and description
        # Try to extract a meaningful description from the first row
        if table.rows:
            first_row_text = " | ".join([cell.text[:15] for cell in table.rows[0].cells])
            caption = f"Table {table_idx}: {first_row_text}..."
        else:
            caption = f"Table {table_idx}"

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.line_spacing = Pt(18)
        p.paragraph_format.left_indent = Inches(0.3)

        rn = p.add_run(caption)
        rn.font.size = Pt(11)
        rn.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

        # Add dots and page number placeholder
        dots = max(2, 55 - len(caption))
        p.add_run(" " + ("." * dots) + " ")
        p.add_run("____").font.size = Pt(10)


def build_list_of_figures_page(doc):
    """Build a List of Figures page (includes charts and images)."""
    add_page_break(doc)

    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run("LIST OF FIGURES")
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    hdr.paragraph_format.space_after = Pt(4)
    add_horizontal_rule(doc, color="1F497D", thickness="8")
    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Count drawing elements (charts/images/figures) in the document
    figure_count = 0
    for element in doc.element.body:
        # Count paragraphs that contain drawing elements
        if element.tag.endswith('}p'):
            for child in element.iter():
                if 'drawing' in child.tag.lower():
                    figure_count += 1
                    break

    if figure_count == 0:
        # No figures found
        p = doc.add_paragraph("No figures in this document.")
        p.paragraph_format.space_after = Pt(6)
        return

    # Generate list of figures
    figure_idx = 0
    for para in doc.paragraphs:
        # Check if paragraph contains images/drawings
        has_drawing = False
        for run in para.runs:
            if run._element.find('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing') is not None:
                has_drawing = True
                break

        if has_drawing:
            figure_idx += 1
            # Try to get caption from next paragraph if it's italicized
            caption = f"Figure {figure_idx}"

            # Look for italicized caption in following paragraphs
            try:
                para_idx = doc.paragraphs.index(para)
                if para_idx + 1 < len(doc.paragraphs):
                    next_para = doc.paragraphs[para_idx + 1]
                    if next_para.runs and next_para.runs[0].italic:
                        caption = f"Figure {figure_idx}: {next_para.text[:50]}"
            except:
                pass

            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(1)
            p.paragraph_format.line_spacing = Pt(18)
            p.paragraph_format.left_indent = Inches(0.3)

            rn = p.add_run(caption)
            rn.font.size = Pt(11)
            rn.font.color.rgb = RGBColor(0x20, 0x20, 0x20)

            # Add dots and page number placeholder
            dots = max(2, 55 - len(caption))
            p.add_run(" " + ("." * dots) + " ")
            p.add_run("____").font.size = Pt(10)


def build_abbreviations_page(doc):
    add_page_break(doc)
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = hdr.add_run("LIST OF ABBREVIATIONS")
    r.font.size  = Pt(14)
    r.font.bold  = True
    r.font.color.rgb = DARK_BLUE
    hdr.paragraph_format.space_after = Pt(4)
    add_horizontal_rule(doc, color="1F497D", thickness="8")

    intro = doc.add_paragraph()
    intro.add_run("The following abbreviations are used throughout this document:").italic = True
    intro.paragraph_format.space_after = Pt(8)

    for abbr, meaning in [
        ("APA",    "American Psychological Association"),
        ("GDP",    "Gross Domestic Product"),
        ("ICT",    "Information and Communication Technology"),
        ("NGO",    "Non-Governmental Organisation"),
        ("UN",     "United Nations"),
        ("WHO",    "World Health Organization"),
        ("SPSS",   "Statistical Package for the Social Sciences"),
        ("e.g.",   "For example (Latin: exempli gratia)"),
        ("i.e.",   "That is (Latin: id est)"),
        ("et al.", "And others (Latin: et alii)"),
    ]:
        p  = doc.add_paragraph(style="List Bullet")
        rb = p.add_run(f"{abbr}:  ")
        rb.bold = True
        p.add_run(meaning)
        p.paragraph_format.space_after = Pt(3)


def build_chapter_page(doc, chapter_num, chapter_content, fn_mgr=None, custom_subtitle=None):
    add_page_break(doc)
    add_chapter_header(doc, chapter_num, custom_subtitle=custom_subtitle)
    parse_chapter_content(doc, chapter_content, fn_mgr=fn_mgr)


def build_document(topic: str, research_level: str,
                   front_matter: str, chapters: dict,
                   output_dir: str) -> str:
    doc = Document()
    set_document_defaults(doc)

    build_title_page(doc, topic, research_level)
    build_front_matter_page(doc, front_matter)
    build_toc_page(doc, research_level)
    build_abbreviations_page(doc)

    # Build all chapters (this creates the tables)
    for num in range(1, 6):
        build_chapter_page(doc, num, chapters[num])

    # Build List of Figures after all chapters (so we can count all figures/charts)
    build_list_of_figures_page(doc)

    # Build List of Tables after all chapters (so we can count all tables)
    build_list_of_tables_page(doc)

    # Sanitize topic for Windows-safe filename
    # Remove or replace invalid Windows filename characters: < > : " / \ | ? *
    safe = topic
    for char in '<>:"/\\|?*':
        safe = safe.replace(char, '_')
    # Keep only alphanumeric, spaces, underscores, and hyphens
    safe = re.sub(r'[^\w\s\-]', '', safe).strip().replace(" ", "_")
    # Limit length and ensure non-empty
    safe = safe[:50] if safe else "Document"

    filename = f"{safe}.docx"
    path     = os.path.join(output_dir, filename)
    doc.save(path)
    return path


# ─────────────────────────────────────────────────────────
#  CLAUDE API — CONTENT GENERATION
# ─────────────────────────────────────────────────────────

def _stream_content(client, system: str, prompt: str,
                    model: str, max_tokens: int, level_key: str = "undergraduate",
                    use_thinking_override: bool = None) -> str:
    """
    Stream content from Claude API with optional extended thinking.

    THINKING ENFORCEMENT:
    - If use_thinking_override=True: ENFORCE thinking (error if model doesn't support)
    - If use_thinking_override=False: Never use thinking, regardless of level
    - If use_thinking_override=None: Use thinking for PhD level IF model supports it
    """
    THINKING_BUDGET = 8000   # tokens reserved for Claude's internal reasoning
    MIN_OUTPUT      = 12000  # minimum tokens guaranteed for actual text output
    THINKING_MODELS = ("claude-opus-4-6", "claude-sonnet-4-6")

    # Determine if thinking should be used
    if use_thinking_override is True:
        # ENFORCE thinking when explicitly requested
        if model not in THINKING_MODELS:
            raise ValueError(
                f"Extended thinking is requested but model '{model}' does not support it. "
                f"Supported models: {', '.join(THINKING_MODELS)}. "
                f"Please use a model that supports extended thinking or disable the thinking toggle."
            )
        use_thinking = True
        print(f"  [THINKING ENABLED] Using extended thinking for deeper analysis", flush=True)
    elif use_thinking_override is False:
        # Explicitly disabled
        use_thinking = False
    else:
        # use_thinking_override is None — use level-based default
        # Thinking for PhD level IF model supports it
        use_thinking = (level_key == "phd") and (model in THINKING_MODELS)

    # Build API kwargs based on thinking setting
    if use_thinking:
        # max_tokens must cover thinking budget + output budget
        # Never let adaptive thinking swallow the output capacity
        actual_max = max(max_tokens, THINKING_BUDGET + MIN_OUTPUT)
        kwargs = dict(
            model=model,
            max_tokens=actual_max,
            thinking={"type": "enabled", "budget_tokens": THINKING_BUDGET},
            system=system,
            messages=[{"role": "user", "content": prompt}],
        )
    else:
        kwargs = dict(
            model=model,
            max_tokens=max_tokens,
            system=system,
            messages=[{"role": "user", "content": prompt}],
        )

    with client.messages.stream(**kwargs) as stream:
        return "".join(stream.text_stream)


def generate_front_matter(client, topic: str, research_level: str,
                           model: str = None,
                           front_matter_sections: list = None,
                           custom_instructions: str = None,
                           use_thinking: bool = False,
                           nalt_compliance: bool = False) -> str:
    """
    Generate front matter pages.

    front_matter_sections : list of optional sections to include.
        Allowed values: "abstract", "declaration", "dedication", "acknowledgements"
        Default (None) → abstract included, other sections excluded.
        Pass ["abstract"] to include only the abstract.
    nalt_compliance : bool, default False
        When True (for BSc legal research), enforce Nigerian Association of
        Law Teachers (NALT) Uniform Format and Citation Guide standards.
    """
    model   = model or config.MODEL
    profile = LEVEL_PROFILES[research_level]

    # Resolve which optional sections to include
    optional_all = ["abstract", "declaration", "dedication", "acknowledgements"]
    if front_matter_sections is None:
        include = ["abstract"]  # Default: only abstract included
    else:
        include = [s.lower().strip() for s in front_matter_sections
                   if s.lower().strip() in optional_all]

    system = (
        f"You are a senior academic writer working at {profile['label']} level. "
        "Write formal, substantive academic front matter. Use ## for each section heading. "
        "Write with authenticity and warmth where appropriate — these pages are read by "
        "examiners and set the tone for the entire document. "
        + HUMAN_WRITING_INSTRUCTION
    )

    if nalt_compliance:
        system += (
            "\n\n"
            "╔═════════════════════════════════════════════════════════════════════╗\n"
            "║  NALT COMPLIANCE MODE: Nigerian Legal Research Standards           ║\n"
            "║  Nigerian Association of Law Teachers (NALT)                       ║\n"
            "║  Uniform Format and Citation Guide for Legal Research Writing      ║\n"
            "╚═════════════════════════════════════════════════════════════════════╝\n\n"
            "STRICT REQUIREMENTS FOR NALT COMPLIANCE:\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "1. FORMATTING (Physical & Technical)\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   • Font: Times New Roman, 12 pt throughout (footnotes auto-reduce to 10pt)\n"
            "   • Line spacing: Double spacing in body; single spacing in footnotes & abstract\n"
            "   • Paper: A4 (210 × 297mm), one-sided only\n"
            "   • Margins: Standard (top/bottom/left/right 2.5cm minimum)\n"
            "   • Page numbers: Bottom right, starting after preliminaries\n"
            "   • Indentation: Indented paragraphs (0.5-1 inch), NOT block text\n"
            "   • Spacing: Single blank line between paragraphs\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "2. DOCUMENT STRUCTURE & CHAPTERS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   FOR BSc LEGAL RESEARCH:\n"
            "   • Total: 50-100 pages (body only)\n"
            "   • Five (5) main chapters:\n"
            "     Ch 1: INTRODUCTION (background, problem statement, aims, objectives, scope)\n"
            "     Ch 2: LITERATURE REVIEW (doctrinal/theoretical foundations)\n"
            "     Ch 3: RESEARCH METHODOLOGY (approach, data collection, analysis)\n"
            "     Ch 4: RESULTS/FINDINGS & DISCUSSION (presentation, interpretation)\n"
            "     Ch 5: CONCLUSIONS & RECOMMENDATIONS\n"
            "   • Preliminaries: Title page, declaration, dedication, acknowledgements, ToC\n"
            "   • Back matter: Bibliography, appendices (if any)\n\n"
            "   HEADINGS HIERARCHY (STRICT):\n"
            "   • Chapter Title: CHAPTER ONE: INTRODUCTION (no number)\n"
            "   • Main heading: 1. Introduction (numbered, ALL CAPS)\n"
            "   • Sub-heading: 1.1 Background to the Study (Title Case)\n"
            "   • Sub-sub-heading: 1.1.1 The Nigerian Context (Title Case, if needed)\n"
            "   DO NOT deviate from this structure\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "3. CITATION FORMAT (FOOTNOTES ONLY - MANDATORY)\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   NEVER use parenthetical citations (not like APA, MLA). ALL citations in footnotes.\n"
            "   Superscript Arabic numerals in text (1, 2, 3...) correspond to footnotes.\n\n"
            "   NIGERIAN CASE LAW FORMAT:\n"
            "   Full citation (first mention):\n"
            "   Okoye v Lagos State Government [1990] 3 NWLR (Pt 136) 115\n"
            "   Short citation (subsequent):\n"
            "   Okoye's case; Okoye case\n"
            "   Unreported cases:\n"
            "   Suit No. ABC/2020 (Lagos High Court)\n\n"
            "   STATUTES & LEGISLATION:\n"
            "   The Evidence Act, CAP E14, Laws of the Federation 2004, section 1\n"
            "   Nigerian Constitution, 1999 (as amended), section 14\n"
            "   Include: Cap letter/number, statute name, year, section cited\n\n"
            "   BOOKS & SECONDARY SOURCES:\n"
            "   First mention: Author name, Title in Single Quotation Marks (Publisher Year) page\n"
            "   Example: Eze Eze, Introduction to Nigerian Constitutional Law (NIALS 2015) 45\n"
            "   Short form: Author surname, Short Title page (for 3+ citations)\n\n"
            "   JOURNAL ARTICLES:\n"
            "   Author, 'Title of Article' (Year) vol/issue page\n"
            "   Example: Dayo Oludele, 'Judicial Activism in Nigeria' (2018) 5(1) NIALS Law Review 23\n\n"
            "   DO NOT include URLs in footnotes; provide page numbers instead\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "4. BIBLIOGRAPHY (END OF DOCUMENT)\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   Arrange in FOUR CATEGORIES (in this order):\n\n"
            "   A. BOOKS (alphabetical by author surname)\n"
            "      Author surname, Initial(s), Title (Edition if applicable) (Publisher Year)\n\n"
            "   B. JOURNAL ARTICLES (alphabetical)\n"
            "      Author, 'Title of Article' (Year) Volume(Issue) Journal-Name page range\n\n"
            "   C. CASES (alphabetical)\n"
            "      Case name [Year] Citation page (if unreported: note court and date)\n\n"
            "   D. STATUTES & LEGISLATIVE INSTRUMENTS (alphabetical)\n"
            "      Statute Name, CAP letter, Laws of the Federation year\n\n"
            "   Within EACH category: strict alphabetical order by surname\n"
            "   When author has multiple works: list in chronological order\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "5. ABSTRACT & FRONT MATTER\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   ABSTRACT:\n"
            "   • Word count: 200-350 words (BSc), 400-500 words (Masters)\n"
            "   • Single spacing, same font as body\n"
            "   • Structure: Background → Problem → Methodology → Key Findings → Conclusion\n"
            "   • KEYWORDS: Maximum 5 keywords, alphabetically ordered, separated by semicolons\n"
            "   • Non-paragraphed (no indentation)\n\n"
            "   PRELIMINARIES (in order):\n"
            "   - Title page\n"
            "   - Declaration of authenticity\n"
            "   - Dedication (if applicable)\n"
            "   - Acknowledgements\n"
            "   - Table of Contents\n"
            "   - Abstract with Keywords\n"
            "   - List of Abbreviations (if any)\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "6. QUOTATIONS & QUOTATION MARKS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   SHORT QUOTATIONS (less than 3 lines / 40 words):\n"
            "   • Integrated into paragraph text\n"
            "   • Use single quotation marks: 'quoted text'\n"
            "   • Cite in footnote immediately after\n"
            "   • Double spacing maintained\n\n"
            "   LONG QUOTATIONS (3+ lines / 40+ words):\n"
            "   • Extract as separate indented block\n"
            "   • Single spacing (not double)\n"
            "   • No quotation marks needed\n"
            "   • Indent 1 inch from both left and right margins\n"
            "   • Blank line before and after quotation\n"
            "   • Cite in footnote immediately after\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "7. RESEARCH APPROACH & CRITICAL ANALYSIS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   DOCTRINAL APPROACH (traditional legal research):\n"
            "   • Textual analysis of statutes and case law\n"
            "   • Analyse Nigerian cases extensively (not just comparative law)\n"
            "   • Trace development of legal principles in Nigerian courts\n"
            "   • Discuss case ratios, distinguishing principles, obiter dicta\n"
            "   • Evaluate statutory provisions with judicial interpretation\n"
            "   • Identify gaps, inconsistencies, ambiguities in Nigerian law\n\n"
            "   NON-DOCTRINAL APPROACH (empirical, socio-legal, etc.):\n"
            "   • Clearly state methodology in Chapter 3\n"
            "   • Specify data collection method (surveys, interviews, observation)\n"
            "   • Define sample, population, data analysis technique\n"
            "   • Present findings with data/evidence\n"
            "   • Relate findings back to legal framework\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "8. LEGAL AUTHORITIES & SOURCE PRIORITIZATION\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   PRIMARY AUTHORITIES (cite extensively):\n"
            "   • Nigerian Constitution, 1999 (as amended)\n"
            "   • Nigerian statutes and legislation\n"
            "   • Nigerian case law (all courts: Supreme Court, Court of Appeal, High Courts)\n"
            "   • Statutory instruments & regulations\n"
            "   • International treaties to which Nigeria is signatory\n\n"
            "   SECONDARY AUTHORITIES (Nigerian authors):\n"
            "   • Nigerian legal scholars' books and journal articles\n"
            "   • NIALS publications\n"
            "   • University law faculty publications\n"
            "   • NALT conference papers\n\n"
            "   COMPARATIVE LAW (use sparingly):\n"
            "   • UK case law only where relevant to Nigerian precedent\n"
            "   • Commonwealth authorities only when applicable\n"
            "   • Always explain why comparative law is cited\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "9. LANGUAGE & ACADEMIC REGISTER\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   • British English spelling (colour, organisation, analyse)\n"
            "   • Formal academic legal register throughout\n"
            "   • Use precise legal terminology consistently\n"
            "   • Avoid jargon, slang, colloquialisms\n"
            "   • Write in third person (avoid 'I', 'we', 'the author')\n"
            "   • Structure arguments in clear prose (not bullet points in body)\n"
            "   • Use transitions and logical connectors between sections\n\n"
            "╚═════════════════════════════════════════════════════════════════════╝\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "CRITICAL: HUMAN AUTHENTICITY WITHIN NALT CONSTRAINTS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "While NALT format is RIGID (citations, structure, bibliography), your WRITING STYLE\n"
            "must remain distinctly human. NALT controls the FORM, not the VOICE.\n\n"
            "WITHIN the constraints, you MUST:\n"
            "• Vary sentence length dramatically (5-48 words per paragraph)\n"
            "• Use unexpected vocabulary choices (not the 'safe' word every time)\n"
            "• Structure arguments unevenly (some paragraphs 1 sentence, others 8)\n"
            "• Show genuine intellectual wrestling with ideas (doubts, tensions, revisions)\n"
            "• Use contractions sparingly but naturally where appropriate\n"
            "• Include rhetorical questions, em dashes, interruptions\n"
            "• Cite scholars with surprise or disagreement, not neutral recitation\n"
            "• Develop arguments idiosyncratically, not formulaically\n"
            "• Show the research process (false starts, complications, nuance)\n\n"
            "Example of COMPLIANT but HUMAN legal writing:\n"
            "  'The case law suggests X. But this reading misses something fundamental. What\n"
            "   the courts have actually done, reading Okoye and its progeny carefully, is adopt\n"
            "   a narrower principle — one that Mensah (2015) flagged but never properly named.\n"
            "   This distinction matters. It changes how we interpret section 14 of the Constitution.'\n\n"
            "DO NOT write:\n"
            "  'The case law indicates that the courts have applied principle X as interpreted\n"
            "   by various scholars. This principle is important to section 14 of the Constitution.'\n\n"
            "Your expertise is showing the HUMAN MIND at work within NALT's formal structure.\n"
            "╚═════════════════════════════════════════════════════════════════════╝"
        )

    # Build the section list dynamically
    section_blocks = []

    if "declaration" in include:
        section_blocks.append(
            "## DECLARATION\n"
            "Write a formal, original academic declaration of authorship (120–160 words). "
            "Include: a statement that the work is the researcher's own, that all sources "
            "have been properly cited, that the work has not been submitted elsewhere for "
            "examination, and acknowledgement that plagiarism constitutes academic misconduct. "
            "The tone should be formal and confident, not mechanical."
        )
    if "dedication" in include:
        section_blocks.append(
            "## DEDICATION\n"
            "Write a heartfelt, personal dedication (60–90 words). Dedicate to specific "
            "named people (family members, mentors, or a community). The dedication should "
            "feel genuine — avoid generic phrases. It should be warm, brief, and memorable."
        )
    if "acknowledgements" in include:
        section_blocks.append(
            "## ACKNOWLEDGEMENTS\n"
            "Write genuine, specific acknowledgements (250–320 words). Thank in separate "
            "sentences or short paragraphs: the academic supervisor (by title and role), "
            "the institution and department, research participants (without naming them), "
            "colleagues or peers who provided feedback, family and close supporters. "
            "The tone should be warm and personal, not formulaic. Avoid generic praise — "
            "be specific about what each person contributed."
        )

    if "abstract" in include:
        abstract_word_min = profile['front_words'] // 2
        abstract_word_max = abstract_word_min + 80
        pg_note = (
            "For postgraduate level: include a sentence on epistemological positioning, "
            "the theoretical framework used, and the study's contribution to theory. "
            if research_level == "postgraduate" else ""
        )
        section_blocks.append(
            f"## ABSTRACT\n"
            f"Write a structured abstract of {abstract_word_min}–{abstract_word_max} words covering: "
            f"(1) background and problem statement, (2) research objectives, "
            f"(3) methodology and data collection approach, (4) principal findings, "
            f"(5) conclusions and recommendations. {pg_note}"
            f"End with: Keywords: [5 relevant academic keywords separated by semicolons]."
        )

    prompt = (
        f"For an academic research project titled: **{topic}**\n"
        f"Research level: {profile['label']}\n\n"
        "Write the following front matter sections. Each must be fully developed to the word counts specified. "
        "Use ## to introduce each section heading.\n\n"
        + "\n\n".join(section_blocks)
        + "\n\nStart directly with the first ## heading. Do not add any preamble or introduction."
    )

    if custom_instructions and custom_instructions.strip():
        prompt += (
            f"\n\n--- ADDITIONAL INSTRUCTIONS ---\n"
            f"{custom_instructions.strip()}\n"
            f"--- END ADDITIONAL INSTRUCTIONS ---"
        )

    print("  [Front Matter] generating...", end=" ", flush=True)
    text = _stream_content(client, system, prompt, model, 500, research_level, use_thinking_override=use_thinking)
    print(f"done ({len(text):,} chars)")

    # Filter out unrequested sections from the generated text
    requested_sections = set()
    for section in include:
        if section == "abstract":
            requested_sections.add("ABSTRACT")
        elif section == "declaration":
            requested_sections.add("DECLARATION")
        elif section == "dedication":
            requested_sections.add("DEDICATION")
        elif section == "acknowledgements":
            requested_sections.add("ACKNOWLEDGEMENTS")

    # Split by ## headings and keep only requested ones
    parts = re.split(r"(?m)^(## .+)$", text)
    filtered_parts = [parts[0]]  # Keep any preamble (usually empty)

    i = 1
    while i + 1 < len(parts):
        heading = parts[i].strip()
        body = parts[i + 1].strip()
        heading_upper = re.sub(r"^## ", "", heading).strip().upper()

        # Keep this section only if it's in the requested list
        if heading_upper in requested_sections:
            filtered_parts.append(heading)
            filtered_parts.append("\n" + body)

        i += 2

    text = "".join(filtered_parts)
    return text


def generate_chapter(client, topic: str, chapter_num: int,
                     research_level: str, model: str = None,
                     custom_instructions: str = None,
                     use_thinking: bool = False,
                     nalt_compliance: bool = False,
                     custom_toc: str = None) -> str:
    model    = model or config.MODEL
    prompts  = _chapter_prompts(research_level, custom_toc=custom_toc, nalt_compliance=nalt_compliance)
    prompt   = prompts[chapter_num].format(topic=topic)
    profile  = LEVEL_PROFILES[research_level]
    target   = profile["word_targets"][chapter_num]

    system = (
        f"You are a highly experienced human academic researcher writing at {profile['label']} level. "
        "You have spent years publishing in peer-reviewed journals and supervising postgraduate students. "
        "Your writing is authoritative, specific, and unmistakably human — characterised by varied rhythm, "
        "genuine intellectual engagement, specific citations, occasional hedging, and a distinct scholarly voice. "
        "You never produce AI-sounding text. You write fully developed, substantive prose and never truncate, "
        "summarise, or leave placeholders. You meet every word count target without compromise."
        + HUMAN_WRITING_INSTRUCTION
    )

    if nalt_compliance:
        system += (
            "\n\n"
            "╔═════════════════════════════════════════════════════════════════════╗\n"
            "║  NALT COMPLIANCE MODE: Nigerian Legal Research Standards           ║\n"
            "║  Nigerian Association of Law Teachers (NALT)                       ║\n"
            "║  Uniform Format and Citation Guide for Legal Research Writing      ║\n"
            "╚═════════════════════════════════════════════════════════════════════╝\n\n"
            "STRICT REQUIREMENTS FOR NALT COMPLIANCE:\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "1. FORMATTING (Physical & Technical)\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   • Font: Times New Roman, 12 pt throughout (footnotes auto-reduce to 10pt)\n"
            "   • Line spacing: Double spacing in body; single spacing in footnotes & abstract\n"
            "   • Paper: A4 (210 × 297mm), one-sided only\n"
            "   • Margins: Standard (top/bottom/left/right 2.5cm minimum)\n"
            "   • Page numbers: Bottom right, starting after preliminaries\n"
            "   • Indentation: Indented paragraphs (0.5-1 inch), NOT block text\n"
            "   • Spacing: Single blank line between paragraphs\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "2. DOCUMENT STRUCTURE & CHAPTERS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   FOR BSc LEGAL RESEARCH:\n"
            "   • Total: 50-100 pages (body only)\n"
            "   • Five (5) main chapters:\n"
            "     Ch 1: INTRODUCTION (background, problem statement, aims, objectives, scope)\n"
            "     Ch 2: LITERATURE REVIEW (doctrinal/theoretical foundations)\n"
            "     Ch 3: RESEARCH METHODOLOGY (approach, data collection, analysis)\n"
            "     Ch 4: RESULTS/FINDINGS & DISCUSSION (presentation, interpretation)\n"
            "     Ch 5: CONCLUSIONS & RECOMMENDATIONS\n"
            "   • Preliminaries: Title page, declaration, dedication, acknowledgements, ToC\n"
            "   • Back matter: Bibliography, appendices (if any)\n\n"
            "   HEADINGS HIERARCHY (STRICT):\n"
            "   • Chapter Title: CHAPTER ONE: INTRODUCTION (no number)\n"
            "   • Main heading: 1. Introduction (numbered, ALL CAPS)\n"
            "   • Sub-heading: 1.1 Background to the Study (Title Case)\n"
            "   • Sub-sub-heading: 1.1.1 The Nigerian Context (Title Case, if needed)\n"
            "   DO NOT deviate from this structure\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "3. CITATION FORMAT (FOOTNOTES ONLY - MANDATORY)\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   NEVER use parenthetical citations (not like APA, MLA). ALL citations in footnotes.\n"
            "   Superscript Arabic numerals in text (1, 2, 3...) correspond to footnotes.\n\n"
            "   NIGERIAN CASE LAW FORMAT:\n"
            "   Full citation (first mention):\n"
            "   Okoye v Lagos State Government [1990] 3 NWLR (Pt 136) 115\n"
            "   Short citation (subsequent):\n"
            "   Okoye's case; Okoye case\n"
            "   Unreported cases:\n"
            "   Suit No. ABC/2020 (Lagos High Court)\n\n"
            "   STATUTES & LEGISLATION:\n"
            "   The Evidence Act, CAP E14, Laws of the Federation 2004, section 1\n"
            "   Nigerian Constitution, 1999 (as amended), section 14\n"
            "   Include: Cap letter/number, statute name, year, section cited\n\n"
            "   BOOKS & SECONDARY SOURCES:\n"
            "   First mention: Author name, Title in Single Quotation Marks (Publisher Year) page\n"
            "   Example: Eze Eze, Introduction to Nigerian Constitutional Law (NIALS 2015) 45\n"
            "   Short form: Author surname, Short Title page (for 3+ citations)\n\n"
            "   JOURNAL ARTICLES:\n"
            "   Author, 'Title of Article' (Year) vol/issue page\n"
            "   Example: Dayo Oludele, 'Judicial Activism in Nigeria' (2018) 5(1) NIALS Law Review 23\n\n"
            "   DO NOT include URLs in footnotes; provide page numbers instead\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "4. BIBLIOGRAPHY (END OF DOCUMENT)\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   Arrange in FOUR CATEGORIES (in this order):\n\n"
            "   A. BOOKS (alphabetical by author surname)\n"
            "      Author surname, Initial(s), Title (Edition if applicable) (Publisher Year)\n\n"
            "   B. JOURNAL ARTICLES (alphabetical)\n"
            "      Author, 'Title of Article' (Year) Volume(Issue) Journal-Name page range\n\n"
            "   C. CASES (alphabetical)\n"
            "      Case name [Year] Citation page (if unreported: note court and date)\n\n"
            "   D. STATUTES & LEGISLATIVE INSTRUMENTS (alphabetical)\n"
            "      Statute Name, CAP letter, Laws of the Federation year\n\n"
            "   Within EACH category: strict alphabetical order by surname\n"
            "   When author has multiple works: list in chronological order\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "5. ABSTRACT & FRONT MATTER\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   ABSTRACT:\n"
            "   • Word count: 200-350 words (BSc), 400-500 words (Masters)\n"
            "   • Single spacing, same font as body\n"
            "   • Structure: Background → Problem → Methodology → Key Findings → Conclusion\n"
            "   • KEYWORDS: Maximum 5 keywords, alphabetically ordered, separated by semicolons\n"
            "   • Non-paragraphed (no indentation)\n\n"
            "   PRELIMINARIES (in order):\n"
            "   - Title page\n"
            "   - Declaration of authenticity\n"
            "   - Dedication (if applicable)\n"
            "   - Acknowledgements\n"
            "   - Table of Contents\n"
            "   - Abstract with Keywords\n"
            "   - List of Abbreviations (if any)\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "6. QUOTATIONS & QUOTATION MARKS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   SHORT QUOTATIONS (less than 3 lines / 40 words):\n"
            "   • Integrated into paragraph text\n"
            "   • Use single quotation marks: 'quoted text'\n"
            "   • Cite in footnote immediately after\n"
            "   • Double spacing maintained\n\n"
            "   LONG QUOTATIONS (3+ lines / 40+ words):\n"
            "   • Extract as separate indented block\n"
            "   • Single spacing (not double)\n"
            "   • No quotation marks needed\n"
            "   • Indent 1 inch from both left and right margins\n"
            "   • Blank line before and after quotation\n"
            "   • Cite in footnote immediately after\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "7. RESEARCH APPROACH & CRITICAL ANALYSIS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   DOCTRINAL APPROACH (traditional legal research):\n"
            "   • Textual analysis of statutes and case law\n"
            "   • Analyse Nigerian cases extensively (not just comparative law)\n"
            "   • Trace development of legal principles in Nigerian courts\n"
            "   • Discuss case ratios, distinguishing principles, obiter dicta\n"
            "   • Evaluate statutory provisions with judicial interpretation\n"
            "   • Identify gaps, inconsistencies, ambiguities in Nigerian law\n\n"
            "   NON-DOCTRINAL APPROACH (empirical, socio-legal, etc.):\n"
            "   • Clearly state methodology in Chapter 3\n"
            "   • Specify data collection method (surveys, interviews, observation)\n"
            "   • Define sample, population, data analysis technique\n"
            "   • Present findings with data/evidence\n"
            "   • Relate findings back to legal framework\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "8. LEGAL AUTHORITIES & SOURCE PRIORITIZATION\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   PRIMARY AUTHORITIES (cite extensively):\n"
            "   • Nigerian Constitution, 1999 (as amended)\n"
            "   • Nigerian statutes and legislation\n"
            "   • Nigerian case law (all courts: Supreme Court, Court of Appeal, High Courts)\n"
            "   • Statutory instruments & regulations\n"
            "   • International treaties to which Nigeria is signatory\n\n"
            "   SECONDARY AUTHORITIES (Nigerian authors):\n"
            "   • Nigerian legal scholars' books and journal articles\n"
            "   • NIALS publications\n"
            "   • University law faculty publications\n"
            "   • NALT conference papers\n\n"
            "   COMPARATIVE LAW (use sparingly):\n"
            "   • UK case law only where relevant to Nigerian precedent\n"
            "   • Commonwealth authorities only when applicable\n"
            "   • Always explain why comparative law is cited\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "9. LANGUAGE & ACADEMIC REGISTER\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "   • British English spelling (colour, organisation, analyse)\n"
            "   • Formal academic legal register throughout\n"
            "   • Use precise legal terminology consistently\n"
            "   • Avoid jargon, slang, colloquialisms\n"
            "   • Write in third person (avoid 'I', 'we', 'the author')\n"
            "   • Structure arguments in clear prose (not bullet points in body)\n"
            "   • Use transitions and logical connectors between sections\n\n"
            "╚═════════════════════════════════════════════════════════════════════╝\n\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "CRITICAL: HUMAN AUTHENTICITY WITHIN NALT CONSTRAINTS\n"
            "━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            "While NALT format is RIGID (citations, structure, bibliography), your WRITING STYLE\n"
            "must remain distinctly human. NALT controls the FORM, not the VOICE.\n\n"
            "WITHIN the constraints, you MUST:\n"
            "• Vary sentence length dramatically (5-48 words per paragraph)\n"
            "• Use unexpected vocabulary choices (not the 'safe' word every time)\n"
            "• Structure arguments unevenly (some paragraphs 1 sentence, others 8)\n"
            "• Show genuine intellectual wrestling with ideas (doubts, tensions, revisions)\n"
            "• Use contractions sparingly but naturally where appropriate\n"
            "• Include rhetorical questions, em dashes, interruptions\n"
            "• Cite scholars with surprise or disagreement, not neutral recitation\n"
            "• Develop arguments idiosyncratically, not formulaically\n"
            "• Show the research process (false starts, complications, nuance)\n\n"
            "Example of COMPLIANT but HUMAN legal writing:\n"
            "  'The case law suggests X. But this reading misses something fundamental. What\n"
            "   the courts have actually done, reading Okoye and its progeny carefully, is adopt\n"
            "   a narrower principle — one that Mensah (2015) flagged but never properly named.\n"
            "   This distinction matters. It changes how we interpret section 14 of the Constitution.'\n\n"
            "DO NOT write:\n"
            "  'The case law indicates that the courts have applied principle X as interpreted\n"
            "   by various scholars. This principle is important to section 14 of the Constitution.'\n\n"
            "Your expertise is showing the HUMAN MIND at work within NALT's formal structure.\n"
            "╚═════════════════════════════════════════════════════════════════════╝"
        )

    if custom_instructions and custom_instructions.strip():
        prompt += (
            f"\n\n--- ADDITIONAL INSTRUCTIONS ---\n"
            f"{custom_instructions.strip()}\n"
            f"--- END ADDITIONAL INSTRUCTIONS ---"
        )

    print(f"  [Ch {chapter_num}] {CHAPTER_SUBTITLES[chapter_num]}...", end=" ", flush=True)
    # Allow token budget based on target word count with 800 token minimum buffer
    text = _stream_content(client, system, prompt, model, max(5000, int(target * 2)), research_level, use_thinking_override=use_thinking)
    print(f"done ({len(text):,} chars)")
    return text


# ─────────────────────────────────────────────────────────
#  CLI ENTRY POINT
# ─────────────────────────────────────────────────────────

def main():
    print("=" * 60)
    print("   ACADEMIC RESEARCH WRITEUP AGENT  —  5-CHAPTER FORMAT")
    print("=" * 60)

    topic = input("\nResearch topic:\n> ").strip()
    level = input("Research level (undergraduate / postgraduate):\n> ").strip().lower()
    if level not in LEVEL_PROFILES:
        print("Defaulting to undergraduate.")
        level = "undergraduate"

    os.environ["ANTHROPIC_API_KEY"] = config.ANTHROPIC_API_KEY
    client = anthropic.Anthropic(api_key=config.ANTHROPIC_API_KEY)

    front    = generate_front_matter(client, topic, level)
    chapters = {n: generate_chapter(client, topic, n, level) for n in range(1, 6)}

    out = build_document(topic, level, front, chapters, os.getcwd())
    print(f"\nDocument saved: {out}")


if __name__ == "__main__":
    main()
